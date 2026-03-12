"""
Report generation service — premium Apple-style design for WorkScanAI
"""
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, PageBreak, HRFlowable, KeepTogether)
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas as rl_canvas
from datetime import datetime
from typing import Dict, List
import os

BLUE        = colors.HexColor('#0071e3')
BLUE_LIGHT  = colors.HexColor('#e8f1fc')
GRAY_900    = colors.HexColor('#1d1d1f')
GRAY_600    = colors.HexColor('#6e6e73')
GRAY_200    = colors.HexColor('#d2d2d7')
GRAY_100    = colors.HexColor('#f5f5f7')
WHITE       = colors.white
GREEN       = colors.HexColor('#34c759')
GREEN_LIGHT = colors.HexColor('#e8f9ed')
AMBER       = colors.HexColor('#ff9f0a')
AMBER_LIGHT = colors.HexColor('#fff4e0')
RED         = colors.HexColor('#ff3b30')
RED_LIGHT   = colors.HexColor('#ffe5e3')


def score_color(score):
    if score >= 70:
        return GREEN, GREEN_LIGHT
    elif score >= 40:
        return AMBER, AMBER_LIGHT
    else:
        return RED, RED_LIGHT


def score_label(score):
    if score >= 70:
        return 'HIGH'
    elif score >= 40:
        return 'MEDIUM'
    else:
        return 'LOW'


class NumberedCanvas(rl_canvas.Canvas):
    def __init__(self, *args, **kwargs):
        rl_canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_footer(num_pages)
            rl_canvas.Canvas.showPage(self)
        rl_canvas.Canvas.save(self)

    def draw_page_footer(self, page_count):
        self.saveState()
        w, h = A4
        self.setStrokeColor(GRAY_200)
        self.setLineWidth(0.5)
        self.line(18*mm, 16*mm, w - 18*mm, 16*mm)
        self.setFont('Helvetica', 8)
        self.setFillColor(GRAY_600)
        self.drawString(18*mm, 11*mm, 'WorkScanAI — AI-Powered Workflow Analysis')
        self.drawRightString(w - 18*mm, 11*mm, f'Page {self._pageNumber} of {page_count}')
        self.restoreState()


class ReportGenerator:

    @staticmethod
    def generate_pdf_report(analysis_data: Dict, output_path: str):
        doc = SimpleDocTemplate(output_path, pagesize=A4,
            leftMargin=18*mm, rightMargin=18*mm, topMargin=20*mm, bottomMargin=24*mm)

        W = A4[0] - 36*mm
        story = []
        s = getSampleStyleSheet()

        def style(name, parent='Normal', **kw):
            return ParagraphStyle(name, parent=s[parent], **kw)

        ST = {
            'cover_title': style('ct', fontSize=22, leading=27, textColor=GRAY_900,
                                  spaceAfter=6, fontName='Helvetica-Bold'),
            'cover_sub':   style('cs', fontSize=13, leading=18, textColor=GRAY_600,
                                  spaceAfter=4, fontName='Helvetica'),
            'cover_meta':  style('cm', fontSize=10, leading=14, textColor=GRAY_600,
                                  fontName='Helvetica'),
            'section_title': style('st', fontSize=18, leading=22, textColor=GRAY_900,
                                    spaceBefore=18, spaceAfter=8, fontName='Helvetica-Bold'),
            'label':  style('lb', fontSize=9, leading=12, textColor=GRAY_600,
                             fontName='Helvetica-Bold', spaceAfter=2),
            'body':   style('bo', fontSize=10, leading=15, textColor=GRAY_900,
                             fontName='Helvetica', spaceAfter=6),
            'rec':    style('rc', fontSize=10, leading=15, textColor=GRAY_900,
                             fontName='Helvetica', leftIndent=10, spaceAfter=8),
        }

        workflow      = analysis_data['workflow']
        score         = analysis_data['automation_score']
        hours         = analysis_data['hours_saved']
        savings       = analysis_data['annual_savings']
        results       = analysis_data['results']
        sorted_results = sorted(results, key=lambda x: x['ai_readiness_score'], reverse=True)
        high = [r for r in results if r['ai_readiness_score'] >= 70]
        med  = [r for r in results if 40 <= r['ai_readiness_score'] < 70]
        low  = [r for r in results if r['ai_readiness_score'] < 40]
        sc, sc_light = score_color(score)

        # ── Cover page ────────────────────────────────────────────────────
        story.append(Table([['']],colWidths=[W],rowHeights=[4],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE)])))
        story.append(Spacer(1, 16*mm))
        story.append(Paragraph('WorkScanAI', style('brand',fontSize=11,textColor=BLUE,fontName='Helvetica-Bold',spaceAfter=10)))
        story.append(Paragraph('Workflow Automation Analysis Report', ST['cover_title']))
        story.append(Paragraph(workflow['name'], style('wn',fontSize=22,leading=26,textColor=BLUE,fontName='Helvetica-Bold',spaceAfter=10)))
        if workflow.get('description'):
            story.append(Paragraph(workflow['description'], ST['cover_sub']))

        # ── Source text right after the title ────────────────────────────
        source_text = workflow.get('source_text', '').strip()
        input_mode  = workflow.get('input_mode', 'manual')
        if source_text:
            mode_labels = {'voice': 'Voice Transcript', 'document': 'Extracted Document Text', 'manual': 'Original Input'}
            mode_label  = mode_labels.get(input_mode, 'Original Input')
            story.append(Spacer(1, 4*mm))
            story.append(Paragraph(mode_label, style('src_lbl', fontSize=9, leading=12, textColor=BLUE,
                fontName='Helvetica-Bold', spaceAfter=4)))
            story.append(Table([[Paragraph(source_text,
                style('src', fontSize=9, leading=14, textColor=GRAY_900, fontName='Helvetica'))]],
                colWidths=[W],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),GRAY_100),
                    ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                    ('LEFTPADDING',(0,0),(-1,-1),14),('RIGHTPADDING',(0,0),(-1,-1),14),
                    ('LINEBEFORE',(0,0),(0,-1),3,BLUE)])))

        story.append(Spacer(1, 8*mm))
        story.append(Paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')}", ST['cover_meta']))
        story.append(Spacer(1, 14*mm))

        hero = Table([[
            Paragraph(f'<font size="48"><b>{score:.0f}%</b></font><br/><font size="14" color="#6e6e73">Automation Score</font>',
                       style('hs',fontName='Helvetica-Bold',alignment=TA_CENTER,leading=54)),
            Table([[Paragraph('<b>Annual Savings</b>',ST['label']),''],
                   [Paragraph(f'\u20ac{savings:,.0f}',style('sv',fontSize=22,fontName='Helvetica-Bold',textColor=GRAY_900)),''],
                   [Paragraph(' ',style('sp1',fontSize=6,fontName='Helvetica')),''],
                   [Paragraph(f'{hours:.0f} hours reclaimed',style('hr',fontSize=10,textColor=GRAY_600,fontName='Helvetica')),'']],
                  colWidths=[W*0.35,W*0.05]),
            Table([[Paragraph('<b>Tasks Analyzed</b>',ST['label']),''],
                   [Paragraph(str(len(results)),style('tc',fontSize=22,fontName='Helvetica-Bold',textColor=GRAY_900)),''],
                   [Paragraph(' ',style('sp2',fontSize=6,fontName='Helvetica')),''],
                   [Paragraph('workflow tasks',style('tl',fontSize=10,textColor=GRAY_600,fontName='Helvetica')),'']],
                  colWidths=[W*0.25,W*0.05]),
        ]],colWidths=[W*0.28,W*0.42,W*0.30],
        style=TableStyle([('BACKGROUND',(0,0),(-1,-1),sc_light),
            ('TOPPADDING',(0,0),(-1,-1),16),('BOTTOMPADDING',(0,0),(-1,-1),16),
            ('LEFTPADDING',(0,0),(0,-1),20),('LEFTPADDING',(1,0),(1,-1),10),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEAFTER',(0,0),(0,-1),1,sc)]))
        story.append(hero)
        story.append(Spacer(1, 10*mm))

        story.append(Table([[
            Paragraph(f'<font color="#34c759"><b>{len(high)} HIGH</b></font> potential', ST['body']),
            Paragraph(f'<font color="#ff9f0a"><b>{len(med)} MEDIUM</b></font> potential', ST['body']),
            Paragraph(f'<font color="#ff3b30"><b>{len(low)} LOW</b></font> potential', ST['body']),
        ]],colWidths=[W/3,W/3,W/3],
        style=TableStyle([('BACKGROUND',(0,0),(-1,-1),GRAY_100),
            ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
            ('LEFTPADDING',(0,0),(-1,-1),14),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('LINEAFTER',(0,0),(1,-1),0.5,GRAY_200)])))
        story.append(PageBreak())

        # ── Detailed Task Analysis ────────────────────────────────────────
        story.append(Paragraph('Detailed Task Analysis', ST['section_title']))
        story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=10))

        for idx, result in enumerate(sorted_results, 1):
            task = result['task']
            task_score = result['ai_readiness_score']
            tc, tc_light = score_color(task_score)
            lbl = score_label(task_score)
            block = []

            header_row = Table([[
                Paragraph(f'<font color="#0071e3"><b>{idx}</b></font>',
                           style(f'tn{idx}',fontSize=22,fontName='Helvetica-Bold',textColor=BLUE,alignment=TA_CENTER)),
                Paragraph(f'<b>{task["name"]}</b>',
                           style(f'th{idx}',fontSize=13,fontName='Helvetica-Bold',textColor=GRAY_900,leading=17)),
                Paragraph(f'<b>{task_score:.0f}%</b><br/><font size="8">{lbl}</font>',
                           style(f'tb{idx}',fontSize=16,fontName='Helvetica-Bold',textColor=tc,alignment=TA_CENTER,leading=20)),
            ]],colWidths=[14*mm,W-14*mm-22*mm,22*mm],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),tc_light),
                ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                ('LEFTPADDING',(0,0),(0,-1),10),('LEFTPADDING',(1,0),(1,-1),8),
                ('RIGHTPADDING',(2,0),(2,-1),10),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                ('LINEBELOW',(0,0),(-1,-1),1.5,tc)]))
            block.append(header_row)

            freq     = task.get('frequency','N/A').capitalize()
            time_pt  = task.get('time_per_task',0)
            cat      = task.get('category','N/A').replace('_',' ').title()
            diff     = result.get('difficulty','N/A').title()
            hrs_sv   = result.get('estimated_hours_saved',0)
            hourly   = analysis_data.get('hourly_rate',50)
            val_sv   = hrs_sv * hourly
            details  = [
                ['Description',   task.get('description','-')],
                ['Frequency',     freq],
                ['Time per Task', f'{time_pt} min'],
                ['Category',      cat],
                ['Implementation',diff],
                ['Annual Savings', f'{hrs_sv:.1f} hrs  /  \u20ac{val_sv:,.0f}'],
            ]
            det_rows = [[Paragraph(f'<b>{k}</b>',ST['label']), Paragraph(v,ST['body'])] for k,v in details]
            block.append(Table(det_rows,colWidths=[42*mm,W-42*mm],
                style=TableStyle([('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('RIGHTPADDING',(0,0),(-1,-1),10),
                    ('VALIGN',(0,0),(-1,-1),'TOP'),
                    ('ROWBACKGROUNDS',(0,0),(-1,-1),[WHITE,GRAY_100]),
                    ('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))

            if result.get('recommendation'):
                import re as _re
                rec_text = result['recommendation']
                opt1_match = _re.search(r'(Option\s+1\s*[—\-–])', rec_text)
                opt2_match = _re.search(r'(Option\s+2\s*[—\-–])', rec_text)
                if opt1_match and opt2_match:
                    before = rec_text[:opt1_match.start()].strip()
                    part1  = rec_text[opt1_match.start():opt2_match.start()].strip()
                    part2  = rec_text[opt2_match.start():].strip()
                    combined = ''
                    if before:
                        combined = f'{before}<br/><br/>'
                    combined += f'{part1}<br/><br/>{part2}'
                    rec_content = Paragraph(combined,
                        style('rec_opt', fontSize=11, fontName='Helvetica', textColor=GRAY_900,
                              leading=16, spaceAfter=0))
                elif opt2_match:
                    part1 = rec_text[:opt2_match.start()].strip()
                    part2 = rec_text[opt2_match.start():].strip()
                    rec_content = Paragraph(
                        f'{part1}<br/><br/>{part2}',
                        style('rec_opt2', fontSize=11, fontName='Helvetica', textColor=GRAY_900,
                              leading=16, spaceAfter=0))
                else:
                    rec_content = Paragraph(rec_text,
                        style('rec_plain', fontSize=11, fontName='Helvetica', textColor=GRAY_900,
                              leading=16, spaceAfter=0))
                block.append(Table([[
                    Paragraph('<b>Recommendation</b>', style('reclbl', fontSize=11,
                        fontName='Helvetica-Bold', textColor=BLUE, leading=14)),
                    rec_content,
                ]],colWidths=[42*mm,W-42*mm],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE_LIGHT),
                    ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('VALIGN',(0,0),(-1,-1),'TOP'),
                    ('LINEBEFORE',(0,0),(0,-1),3,BLUE)])))

            block.append(Spacer(1,8*mm))
            story.append(KeepTogether(block))

        # ── Roadmap ───────────────────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Implementation Roadmap', ST['section_title']))
        story.append(HRFlowable(width=W,thickness=0.5,color=GRAY_200,spaceAfter=10))

        phases = [
            ('Phase 1 — Quick Wins','0–3 months','High score + easy setup = immediate ROI.',
             GREEN, GREEN_LIGHT,
             [r for r in sorted_results if r['ai_readiness_score']>=70 and r.get('difficulty','').lower()=='easy']),
            ('Phase 2 — Medium-Term','3–6 months','Technical setup required but significant savings.',
             AMBER, AMBER_LIGHT,
             [r for r in sorted_results if r['ai_readiness_score']>=50 and r.get('difficulty','').lower()=='medium']),
            ('Phase 3 — Advanced','6–12 months','Complex automations and custom development.',
             BLUE, BLUE_LIGHT,
             [r for r in sorted_results if r['ai_readiness_score']>=40 and r.get('difficulty','').lower()=='hard']),
        ]
        for title,timeline,desc,col,col_light,phase_tasks in phases:
            pb = []
            pb.append(Table([[
                Paragraph(f'<b>{title}</b>',style(f'ph{title}',fontSize=13,fontName='Helvetica-Bold',textColor=col)),
                Paragraph(timeline,style(f'pt{title}',fontSize=10,textColor=GRAY_600,fontName='Helvetica',alignment=TA_RIGHT)),
            ]],colWidths=[W*0.65,W*0.35],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),col_light),
                ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                ('LEFTPADDING',(0,0),(0,-1),12),('RIGHTPADDING',(-1,0),(-1,-1),12),
                ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEBELOW',(0,0),(-1,-1),1.5,col)])))
            pb.append(Paragraph(desc,style(f'pd{title}',fontSize=9,textColor=GRAY_600,fontName='Helvetica',spaceAfter=6,leftIndent=12)))
            if phase_tasks:
                for r in phase_tasks:
                    pb.append(Table([[
                        Paragraph('>',style(f'ar{r["task"]["name"]}',fontSize=10,fontName='Helvetica-Bold',textColor=col)),
                        Paragraph(f'<b>{r["task"]["name"]}</b>  <font size="9" color="#6e6e73">Score {r["ai_readiness_score"]:.0f}%  {r.get("estimated_hours_saved",0):.0f} hrs saved</font>',
                                   style(f'ri{r["task"]["name"]}',fontSize=10,fontName='Helvetica',textColor=GRAY_900,leading=14)),
                    ]],colWidths=[8*mm,W-8*mm],
                    style=TableStyle([('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
                        ('LEFTPADDING',(0,0),(-1,-1),12),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            else:
                pb.append(Paragraph('No tasks in this phase.',style(f'np{title}',fontSize=9,textColor=GRAY_600,fontName='Helvetica-Oblique',leftIndent=12,spaceAfter=4)))
            pb.append(Spacer(1,8*mm))
            story.append(KeepTogether(pb))

        # ── Conclusion ────────────────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Conclusion', ST['section_title']))
        story.append(HRFlowable(width=W,thickness=0.5,color=GRAY_200,spaceAfter=10))
        story.append(Paragraph(
            f'This analysis identified automation opportunities across <b>{len(results)} tasks</b> '
            f'in <b>{workflow["name"]}</b>. Implementing recommendations could save '
            f'<b>{hours:.0f} hours annually</b>, worth approximately <b>\u20ac{savings:,.0f}</b>.',
            ST['body']))
        story.append(Spacer(1,6*mm))
        for item in [
            'Begin with Phase 1 quick wins — ROI within weeks, not months.',
            'Use human-in-the-loop for tasks scored 40–70 to reduce risk.',
            'Track time saved monthly to measure and communicate impact.',
            'Revisit quarterly — new AI tools appear constantly.',
            'Train team members on newly automated workflows.',
        ]:
            story.append(Table([[
                Paragraph('>',style(f'ai{item[:10]}',fontSize=11,textColor=BLUE,fontName='Helvetica-Bold')),
                Paragraph(item,ST['body']),
            ]],colWidths=[8*mm,W-8*mm],
            style=TableStyle([('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
                ('LEFTPADDING',(0,0),(0,-1),0),('VALIGN',(0,0),(-1,-1),'TOP')])))

        doc.build(story, canvasmaker=NumberedCanvas)
        return output_path

    # ── DOCX ─────────────────────────────────────────────────────────────────
    @staticmethod
    def generate_docx_report(analysis_data: Dict, output_path: str):
        if Document is None:
            raise ImportError("python-docx not installed")
        doc = Document()
        for section in doc.sections:
            section.top_margin=Cm(2); section.bottom_margin=Cm(2)
            section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)

        workflow=analysis_data['workflow']; score=analysis_data['automation_score']
        hours=analysis_data['hours_saved']; savings=analysis_data['annual_savings']
        results=analysis_data['results']
        sorted_results=sorted(results,key=lambda x:x['ai_readiness_score'],reverse=True)

        def set_cell_bg(cell,hex_color):
            tc=cell._tc; tcPr=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
            tcPr.append(shd)

        def set_col_width(cell,width_cm):
            tc=cell._tc; tcPr=tc.get_or_add_tcPr()
            tcW=OxmlElement('w:tcW')
            tcW.set(qn('w:w'),str(int(width_cm*567))); tcW.set(qn('w:type'),'dxa')
            tcPr.append(tcW)

        def add_run(para,text,bold=False,size=None,color=None,italic=False):
            run=para.add_run(text); run.bold=bold; run.italic=italic
            if size: run.font.size=Pt(size)
            if color: run.font.color.rgb=RGBColor(*tuple(int(color[i:i+2],16) for i in (0,2,4)))
            return run

        def add_section_heading(text):
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
            run=p.add_run(text); run.bold=True; run.font.size=Pt(16); run.font.color.rgb=RGBColor(0x1d,0x1d,0x1f)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
            bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'4')
            bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'d2d2d7')
            pBdr.append(bottom); pPr.append(pBdr)

        # Cover
        p=doc.add_paragraph(); add_run(p,'WorkScanAI',bold=True,size=11,color='0071e3')
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); add_run(p,'Workflow Automation Analysis Report',bold=True,size=22,color='1d1d1f')
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
        add_run(p,workflow['name'],bold=True,size=20,color='0071e3')

        if workflow.get('description'):
            p=doc.add_paragraph(); add_run(p,workflow['description'],size=11,color='6e6e73'); p.paragraph_format.space_after=Pt(8)

        # ── Source text right after the title ────────────────────────────
        source_text = workflow.get('source_text', '').strip()
        input_mode  = workflow.get('input_mode', 'manual')
        if source_text:
            mode_labels = {'voice': 'Voice Transcript', 'document': 'Extracted Document Text', 'manual': 'Original Input'}
            mode_label  = mode_labels.get(input_mode, 'Original Input')
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
            add_run(p, mode_label, bold=True, size=9, color='0071e3')
            st_tbl = doc.add_table(rows=1, cols=1); st_tbl.style = 'Table Grid'
            cell = st_tbl.rows[0].cells[0]
            set_cell_bg(cell, 'f5f5f7')
            p = cell.paragraphs[0]
            add_run(p, source_text, size=10, color='1d1d1f')
            doc.add_paragraph()

        p=doc.add_paragraph(); add_run(p,f"Generated {datetime.now().strftime('%B %d, %Y')}",size=9,color='6e6e73')
        p.paragraph_format.space_after=Pt(12)

        high=[r for r in results if r['ai_readiness_score']>=70]
        med=[r for r in results if 40<=r['ai_readiness_score']<70]
        low_=[r for r in results if r['ai_readiness_score']<40]

        tbl=doc.add_table(rows=2,cols=4); tbl.style='Table Grid'
        headers=['Automation Score','Annual Savings','Hours Reclaimed','Tasks Analyzed']
        values=[f'{score:.0f}%',f'\u20ac{savings:,.0f}',f'{hours:.0f} hrs',str(len(results))]
        bg_colors=['e8f1fc','e8f9ed','fff4e0','f5f5f7']
        for i,(h,v,bg) in enumerate(zip(headers,values,bg_colors)):
            hc=tbl.rows[0].cells[i]; vc=tbl.rows[1].cells[i]
            set_cell_bg(hc,bg); set_cell_bg(vc,bg)
            ph=hc.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(ph,h,bold=True,size=8,color='6e6e73')
            pv=vc.paragraphs[0]; pv.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(pv,v,bold=True,size=16,color='1d1d1f')
        doc.add_paragraph()

        add_section_heading('Task Breakdown')
        for idx, result in enumerate(results, 1):
            task=result['task']; task_score=result['ai_readiness_score']; lbl=score_label(task_score)
            score_hex,bg_hex=('34c759','e8f9ed') if task_score>=70 else (('ff9f0a','fff4e0') if task_score>=40 else ('ff3b30','ffe5e3'))

            tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'
            c0=tbl.rows[0].cells[0]; set_cell_bg(c0,bg_hex); p0=c0.paragraphs[0]
            p0.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(p0,f'{idx:02d}',bold=True,size=18,color='0071e3')
            c1=tbl.rows[0].cells[1]; set_cell_bg(c1,bg_hex); p1=c1.paragraphs[0]
            add_run(p1,task['name'],bold=True,size=13,color='1d1d1f')
            c2=tbl.rows[0].cells[2]; set_cell_bg(c2,bg_hex); p2=c2.paragraphs[0]
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            add_run(p2,f'{task_score:.0f}%\n',bold=True,size=16,color=score_hex)
            add_run(p2,lbl,bold=True,size=8,color=score_hex)
            set_col_width(c0,1.2); set_col_width(c1,11.0); set_col_width(c2,2.8)
            doc.add_paragraph()

            freq=task.get('frequency','N/A').capitalize(); time_pt=task.get('time_per_task',0)
            cat=task.get('category','N/A').replace('_',' ').title(); diff=result.get('difficulty','N/A').title()
            hrs_sv=result.get('estimated_hours_saved',0); hourly=analysis_data.get('hourly_rate',50); val_sv=hrs_sv*hourly
            details=[('Description',task.get('description','-')),('Frequency',freq),
                     ('Time per Task',f'{time_pt} minutes'),('Category',cat),
                     ('Difficulty',diff),('Annual Savings',f'{hrs_sv:.1f} hrs  /  \u20ac{val_sv:,.0f}')]
            dt=doc.add_table(rows=len(details),cols=2); dt.style='Table Grid'
            for row_i,(k,v) in enumerate(details):
                kc=dt.rows[row_i].cells[0]; vc=dt.rows[row_i].cells[1]
                bg='f5f5f7' if row_i%2==0 else 'ffffff'
                set_cell_bg(kc,'f5f5f7'); set_cell_bg(vc,bg)
                set_col_width(kc,4.0); set_col_width(vc,11.0)
                add_run(kc.paragraphs[0],k,bold=True,size=9,color='6e6e73')
                add_run(vc.paragraphs[0],v,size=10,color='1d1d1f')
            if result.get('recommendation'):
                doc.add_paragraph()
                rt=doc.add_table(rows=1,cols=2); rt.style='Table Grid'
                lc=rt.rows[0].cells[0]; rc=rt.rows[0].cells[1]
                set_cell_bg(lc,'e8f1fc'); set_cell_bg(rc,'e8f1fc')
                set_col_width(lc,3.8); set_col_width(rc,11.2)
                add_run(lc.paragraphs[0],'Recommendation',bold=True,size=12,color='0071e3')
                import re as _re2
                rec_text = result['recommendation']
                opt1_m = _re2.search(r'(Option\s+1\s*[—\-–])', rec_text)
                opt2_m = _re2.search(r'(Option\s+2\s*[—\-–])', rec_text)
                if opt1_m and opt2_m:
                    before = rec_text[:opt1_m.start()].strip()
                    part1  = rec_text[opt1_m.start():opt2_m.start()].strip()
                    part2  = rec_text[opt2_m.start():].strip()
                    if before:
                        add_run(rc.paragraphs[0], before, size=11, color='1d1d1f')
                        p_mid = rc.add_paragraph()
                        add_run(p_mid, part1, size=11, color='1d1d1f')
                    else:
                        add_run(rc.paragraphs[0], part1, size=11, color='1d1d1f')
                    p2 = rc.add_paragraph()
                    add_run(p2, part2, size=11, color='1d1d1f')
                elif opt2_m:
                    part1 = rec_text[:opt2_m.start()].strip()
                    part2 = rec_text[opt2_m.start():].strip()
                    add_run(rc.paragraphs[0], part1, size=11, color='1d1d1f')
                    p2 = rc.add_paragraph()
                    add_run(p2, part2, size=11, color='1d1d1f')
                else:
                    add_run(rc.paragraphs[0], rec_text, size=11, color='1d1d1f')
            doc.add_paragraph()

        add_section_heading('Implementation Roadmap')
        for phase_title,phase_tasks,col_hex in [
            ('Phase 1 — Quick Wins (0–3 months)',[r for r in sorted_results if r['ai_readiness_score']>=70 and r.get('difficulty','').lower()=='easy'],'34c759'),
            ('Phase 2 — Medium-Term (3–6 months)',[r for r in sorted_results if r['ai_readiness_score']>=50 and r.get('difficulty','').lower()=='medium'],'ff9f0a'),
            ('Phase 3 — Advanced (6–12 months)',[r for r in sorted_results if r['ai_readiness_score']>=40 and r.get('difficulty','').lower()=='hard'],'0071e3'),
        ]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10)
            add_run(p,phase_title,bold=True,size=12,color=col_hex)
            if phase_tasks:
                for r in phase_tasks:
                    p=doc.add_paragraph(style='List Bullet')
                    add_run(p,r['task']['name'],bold=True,size=10,color='1d1d1f')
                    add_run(p,f"  /  Score {r['ai_readiness_score']:.0f}%  /  {r.get('estimated_hours_saved',0):.0f} hrs saved",size=9,color='6e6e73')
            else:
                p=doc.add_paragraph(); add_run(p,'No tasks in this phase.',size=9,color='6e6e73',italic=True)

        add_section_heading('Conclusion')
        p=doc.add_paragraph()
        add_run(p,f'This analysis identified automation opportunities across {len(results)} tasks in {workflow["name"]}. ',size=10,color='1d1d1f')
        add_run(p,f'Implementing recommendations could save {hours:.0f} hours annually',bold=True,size=10,color='1d1d1f')
        add_run(p,f', worth approximately ',size=10,color='1d1d1f')
        add_run(p,f'\u20ac{savings:,.0f}',bold=True,size=10,color='1d1d1f')
        add_run(p,'.',size=10,color='1d1d1f')
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8)
        add_run(p,'Key Recommendations:',bold=True,size=11,color='1d1d1f')
        for step in ['Start with Phase 1 quick wins for immediate ROI within weeks.',
                     'Apply human-in-the-loop for tasks scored 40-70 to manage risk.',
                     'Track time saved monthly to measure and communicate impact.',
                     'Revisit quarterly — new AI tools appear constantly.',
                     'Train team members on newly automated workflows.']:
            p=doc.add_paragraph(style='List Bullet'); add_run(p,step,size=10,color='1d1d1f')
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        add_run(p,'Generated by WorkScanAI — AI-Powered Workflow Analysis',size=8,color='6e6e73')
        doc.save(output_path)
        return output_path

    # ── Combined multi-workflow reports ───────────────────────────────────────

    @staticmethod
    def generate_combined_pdf_report(analyses_list: List[Dict], output_path: str):
        """One PDF with all workflows — each separated by a cover divider page."""
        doc = SimpleDocTemplate(output_path, pagesize=A4,
            leftMargin=18*mm, rightMargin=18*mm, topMargin=20*mm, bottomMargin=24*mm)
        W = A4[0] - 36*mm
        story = []
        s = getSampleStyleSheet()

        def style(name, parent='Normal', **kw):
            return ParagraphStyle(name, parent=s[parent], **kw)

        ST = {
            'cover_title': style('ct2', fontSize=32, leading=38, textColor=GRAY_900,
                                  spaceAfter=6, fontName='Helvetica-Bold'),
            'section_title': style('st2', fontSize=18, leading=22, textColor=GRAY_900,
                                    spaceBefore=18, spaceAfter=8, fontName='Helvetica-Bold'),
            'label': style('lb2', fontSize=9, leading=12, textColor=GRAY_600,
                            fontName='Helvetica-Bold', spaceAfter=2),
            'body':  style('bo2', fontSize=10, leading=15, textColor=GRAY_900,
                            fontName='Helvetica', spaceAfter=6),
            'rec':   style('rc2', fontSize=10, leading=15, textColor=GRAY_900,
                            fontName='Helvetica', leftIndent=10, spaceAfter=8),
        }

        total_savings = sum(a['annual_savings'] for a in analyses_list)
        total_hours   = sum(a['hours_saved']    for a in analyses_list)
        total_tasks   = sum(len(a['results'])   for a in analyses_list)

        # ── Master cover page ─────────────────────────────────────────────
        story.append(Table([['']],colWidths=[W],rowHeights=[4],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE)])))
        story.append(Spacer(1, 16*mm))
        story.append(Paragraph('WorkScanAI', style('brand2',fontSize=11,textColor=BLUE,fontName='Helvetica-Bold',spaceAfter=10)))
        story.append(Paragraph('Combined Workflow Automation\nAnalysis Report', ST['cover_title']))
        story.append(Paragraph(f'{len(analyses_list)} Workflows', style('wn2',fontSize=22,leading=26,textColor=BLUE,fontName='Helvetica-Bold',spaceAfter=10)))
        story.append(Paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')}", style('cm2',fontSize=10,leading=14,textColor=GRAY_600,fontName='Helvetica')))
        story.append(Spacer(1, 14*mm))

        story.append(Table([[
            Paragraph(f'<font size="36"><b>{len(analyses_list)}</b></font><br/><font size="12" color="#6e6e73">workflows</font>',
                       style('hs2',fontName='Helvetica-Bold',alignment=TA_CENTER,leading=42)),
            Table([[Paragraph('<b>Total Savings</b>',ST['label']),''],
                   [Paragraph(f'\u20ac{total_savings:,.0f}',style('sv2',fontSize=20,fontName='Helvetica-Bold',textColor=GRAY_900)),''],
                   [Paragraph(f'{total_hours:.0f} hours reclaimed',style('hr2',fontSize=10,textColor=GRAY_600,fontName='Helvetica')),'']],
                  colWidths=[W*0.38,W*0.05]),
            Table([[Paragraph('<b>Total Tasks</b>',ST['label']),''],
                   [Paragraph(str(total_tasks),style('tc2',fontSize=20,fontName='Helvetica-Bold',textColor=GRAY_900)),''],
                   [Paragraph('analyzed',style('tl2',fontSize=10,textColor=GRAY_600,fontName='Helvetica')),'']],
                  colWidths=[W*0.25,W*0.05]),
        ]],colWidths=[W*0.27,W*0.44,W*0.29],
        style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE_LIGHT),
            ('TOPPADDING',(0,0),(-1,-1),16),('BOTTOMPADDING',(0,0),(-1,-1),16),
            ('LEFTPADDING',(0,0),(0,-1),20),('LEFTPADDING',(1,0),(1,-1),10),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEAFTER',(0,0),(0,-1),1,BLUE)])))
        story.append(Spacer(1, 10*mm))

        # Index table
        idx_rows = [[Paragraph('<b>#</b>',ST['label']),
                     Paragraph('<b>Workflow</b>',ST['label']),
                     Paragraph('<b>Score</b>',ST['label']),
                     Paragraph('<b>Savings</b>',ST['label'])]]
        for i, a in enumerate(analyses_list, 1):
            sc, _ = score_color(a['automation_score'])
            idx_rows.append([
                Paragraph(str(i), style(f'ix{i}',fontSize=10,fontName='Helvetica',textColor=GRAY_600)),
                Paragraph(a['workflow']['name'], style(f'iw{i}',fontSize=10,fontName='Helvetica-Bold',textColor=GRAY_900)),
                Paragraph(f"{a['automation_score']:.0f}%", style(f'is{i}',fontSize=10,fontName='Helvetica-Bold',textColor=sc)),
                Paragraph(f"\u20ac{a['annual_savings']:,.0f}", style(f'iv{i}',fontSize=10,fontName='Helvetica',textColor=GRAY_900)),
            ])
        story.append(Table(idx_rows, colWidths=[10*mm, W-10*mm-30*mm-40*mm, 30*mm, 40*mm],
            style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                ('BACKGROUND',(0,1),(-1,-1),WHITE),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                ('TOPPADDING',(0,0),(-1,-1),7),('BOTTOMPADDING',(0,0),(-1,-1),7),
                ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
        story.append(PageBreak())

        # ── Per-workflow sections ─────────────────────────────────────────
        for w_idx, analysis_data in enumerate(analyses_list):
            workflow      = analysis_data['workflow']
            score         = analysis_data['automation_score']
            hours         = analysis_data['hours_saved']
            savings       = analysis_data['annual_savings']
            results       = analysis_data['results']
            sorted_results = sorted(results, key=lambda x: x['ai_readiness_score'], reverse=True)
            sc, sc_light  = score_color(score)

            # Workflow divider cover
            story.append(Table([['']],colWidths=[W],rowHeights=[3],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),sc)])))
            story.append(Spacer(1, 10*mm))
            story.append(Paragraph(f'Workflow {w_idx+1} of {len(analyses_list)}',
                style(f'wl{w_idx}',fontSize=10,textColor=GRAY_600,fontName='Helvetica',spaceAfter=6)))
            story.append(Paragraph(workflow['name'],
                style(f'wt{w_idx}',fontSize=26,leading=30,textColor=GRAY_900,fontName='Helvetica-Bold',spaceAfter=8)))
            if workflow.get('description'):
                story.append(Paragraph(workflow['description'],
                    style(f'wd{w_idx}',fontSize=12,textColor=GRAY_600,fontName='Helvetica',spaceAfter=8)))

            story.append(Table([[
                Paragraph(f'<b>{score:.0f}%</b><br/><font size="9" color="#6e6e73">Automation Score</font>',
                           style(f'ws{w_idx}',fontSize=22,fontName='Helvetica-Bold',textColor=sc,alignment=TA_CENTER,leading=28)),
                Paragraph(f'<b>\u20ac{savings:,.0f}</b><br/><font size="9" color="#6e6e73">Annual Savings</font>',
                           style(f'wv{w_idx}',fontSize=18,fontName='Helvetica-Bold',textColor=GRAY_900,alignment=TA_CENTER,leading=24)),
                Paragraph(f'<b>{hours:.0f} hrs</b><br/><font size="9" color="#6e6e73">Per Year Saved</font>',
                           style(f'wh{w_idx}',fontSize=18,fontName='Helvetica-Bold',textColor=GRAY_900,alignment=TA_CENTER,leading=24)),
                Paragraph(f'<b>{len(results)}</b><br/><font size="9" color="#6e6e73">Tasks</font>',
                           style(f'wc{w_idx}',fontSize=18,fontName='Helvetica-Bold',textColor=GRAY_900,alignment=TA_CENTER,leading=24)),
            ]],colWidths=[W/4]*4,
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),sc_light),
                ('TOPPADDING',(0,0),(-1,-1),14),('BOTTOMPADDING',(0,0),(-1,-1),14),
                ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                ('LINEAFTER',(0,0),(2,-1),0.5,sc)])))
            story.append(Spacer(1, 8*mm))

            story.append(Paragraph('Task Analysis', ST['section_title']))
            story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=8))

            for idx, result in enumerate(sorted_results, 1):
                task = result['task']
                task_score = result['ai_readiness_score']
                tc, tc_light = score_color(task_score)
                lbl = score_label(task_score)
                block = []
                uid = f'w{w_idx}t{idx}'

                block.append(Table([[
                    Paragraph(f'<font color="#0071e3"><b>{idx}</b></font>',
                               style(f'tn{uid}',fontSize=18,fontName='Helvetica-Bold',textColor=BLUE,alignment=TA_CENTER)),
                    Paragraph(f'<b>{task["name"]}</b>',
                               style(f'th{uid}',fontSize=12,fontName='Helvetica-Bold',textColor=GRAY_900,leading=16)),
                    Paragraph(f'<b>{task_score:.0f}%</b><br/><font size="8">{lbl}</font>',
                               style(f'tb{uid}',fontSize=14,fontName='Helvetica-Bold',textColor=tc,alignment=TA_CENTER,leading=18)),
                ]],colWidths=[12*mm,W-12*mm-20*mm,20*mm],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),tc_light),
                    ('TOPPADDING',(0,0),(-1,-1),8),('BOTTOMPADDING',(0,0),(-1,-1),8),
                    ('LEFTPADDING',(0,0),(-1,-1),8),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                    ('LINEBELOW',(0,0),(-1,-1),1.5,tc)])))

                hrs_sv = result.get('estimated_hours_saved', 0)
                details = [
                    ['Difficulty', result.get('difficulty','N/A').title()],
                    ['Annual Hours Saved', f'{hrs_sv:.1f} hrs'],
                ]
                block.append(Table(
                    [[Paragraph(f'<b>{k}</b>',ST['label']), Paragraph(v,ST['body'])] for k,v in details],
                    colWidths=[32*mm,W-32*mm],
                    style=TableStyle([('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
                        ('LEFTPADDING',(0,0),(-1,-1),8),('ROWBACKGROUNDS',(0,0),(-1,-1),[WHITE,GRAY_100]),
                        ('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))

                if result.get('recommendation'):
                    block.append(Table([[
                        Paragraph('Recommendation', ST['label']),
                        Paragraph(result['recommendation'], ST['rec']),
                    ]],colWidths=[32*mm,W-32*mm],
                    style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE_LIGHT),
                        ('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6),
                        ('LEFTPADDING',(0,0),(-1,-1),8),('VALIGN',(0,0),(-1,-1),'TOP'),
                        ('LINEBEFORE',(0,0),(0,-1),3,BLUE)])))

                block.append(Spacer(1, 6*mm))
                story.append(KeepTogether(block))

            if w_idx < len(analyses_list) - 1:
                story.append(PageBreak())

        # ── Combined conclusion ───────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Combined Summary', ST['section_title']))
        story.append(HRFlowable(width=W,thickness=0.5,color=GRAY_200,spaceAfter=10))
        story.append(Paragraph(
            f'This combined report covers <b>{len(analyses_list)} workflows</b> with '
            f'<b>{total_tasks} total tasks</b> analyzed. Across all workflows, implementing '
            f'automation recommendations could save <b>{total_hours:.0f} hours annually</b>, '
            f'worth approximately <b>\u20ac{total_savings:,.0f}</b>.',
            style('conc',fontSize=10,leading=15,textColor=GRAY_900,fontName='Helvetica',spaceAfter=6)))

        doc.build(story, canvasmaker=NumberedCanvas)
        return output_path

    @staticmethod
    def generate_combined_docx_report(analyses_list: List[Dict], output_path: str):
        """One DOCX with all workflows separated by divider headings."""
        if Document is None:
            raise ImportError("python-docx not installed")
        doc = Document()
        for section in doc.sections:
            section.top_margin=Cm(2); section.bottom_margin=Cm(2)
            section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)

        def set_cell_bg(cell, hex_color):
            tc=cell._tc; tcPr=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
            tcPr.append(shd)

        def set_col_width(cell, width_cm):
            tc=cell._tc; tcPr=tc.get_or_add_tcPr()
            tcW=OxmlElement('w:tcW')
            tcW.set(qn('w:w'),str(int(width_cm*567))); tcW.set(qn('w:type'),'dxa')
            tcPr.append(tcW)

        def add_run(para, text, bold=False, size=None, color=None, italic=False):
            run=para.add_run(text); run.bold=bold; run.italic=italic
            if size: run.font.size=Pt(size)
            if color: run.font.color.rgb=RGBColor(*tuple(int(color[i:i+2],16) for i in (0,2,4)))
            return run

        def add_heading(text, size=16, color='1d1d1f'):
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
            run=p.add_run(text); run.bold=True; run.font.size=Pt(size)
            run.font.color.rgb=RGBColor(*tuple(int(color[i:i+2],16) for i in (0,2,4)))
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
            bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'4')
            bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'d2d2d7')
            pBdr.append(bottom); pPr.append(pBdr)

        total_savings = sum(a['annual_savings'] for a in analyses_list)
        total_hours   = sum(a['hours_saved']    for a in analyses_list)
        total_tasks   = sum(len(a['results'])   for a in analyses_list)

        # Master cover
        p=doc.add_paragraph(); add_run(p,'WorkScanAI',bold=True,size=11,color='0071e3')
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4)
        add_run(p,'Combined Workflow Automation Analysis Report',bold=True,size=28,color='1d1d1f')
        p=doc.add_paragraph(); add_run(p,f'{len(analyses_list)} Workflows  ·  {total_tasks} Tasks  ·  \u20ac{total_savings:,.0f} Potential Savings',size=13,color='6e6e73')
        p=doc.add_paragraph(); add_run(p,f"Generated {datetime.now().strftime('%B %d, %Y')}",size=9,color='6e6e73')
        p.paragraph_format.space_after=Pt(14)

        # Summary table
        tbl=doc.add_table(rows=2,cols=3); tbl.style='Table Grid'
        for i,(h,v,bg) in enumerate([('Total Savings',f'\u20ac{total_savings:,.0f}','e8f9ed'),
                                      ('Hours Saved',f'{total_hours:.0f} hrs','e8f1fc'),
                                      ('Total Tasks',str(total_tasks),'f5f5f7')]):
            hc=tbl.rows[0].cells[i]; vc=tbl.rows[1].cells[i]
            set_cell_bg(hc,bg); set_cell_bg(vc,bg)
            ph=hc.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(ph,h,bold=True,size=9,color='6e6e73')
            pv=vc.paragraphs[0]; pv.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(pv,v,bold=True,size=18,color='1d1d1f')
        doc.add_paragraph()

        # Per-workflow sections
        for w_idx, analysis_data in enumerate(analyses_list):
            workflow=analysis_data['workflow']; score=analysis_data['automation_score']
            hours=analysis_data['hours_saved']; savings=analysis_data['annual_savings']
            results=analysis_data['results']
            sorted_results=sorted(results,key=lambda x:x['ai_readiness_score'],reverse=True)
            score_hex='34c759' if score>=70 else ('ff9f0a' if score>=40 else 'ff3b30')

            # Workflow divider
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(2)
            add_run(p,f'WORKFLOW {w_idx+1} OF {len(analyses_list)}',bold=True,size=9,color='6e6e73')
            add_heading(workflow['name'], size=20, color='1d1d1f')

            # Workflow stats row
            tbl=doc.add_table(rows=2,cols=4); tbl.style='Table Grid'
            for i,(h,v,bg) in enumerate([('Automation Score',f'{score:.0f}%','e8f1fc'),
                                          ('Annual Savings',f'\u20ac{savings:,.0f}','e8f9ed'),
                                          ('Hours Saved',f'{hours:.0f} hrs','fff4e0'),
                                          ('Tasks',str(len(results)),'f5f5f7')]):
                hc=tbl.rows[0].cells[i]; vc=tbl.rows[1].cells[i]
                set_cell_bg(hc,bg); set_cell_bg(vc,bg)
                ph=hc.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(ph,h,bold=True,size=8,color='6e6e73')
                pv=vc.paragraphs[0]; pv.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(pv,v,bold=True,size=14,color='1d1d1f')
            doc.add_paragraph()

            add_heading('Task Analysis', size=14)
            for idx,result in enumerate(sorted_results,1):
                task=result['task']; task_score=result['ai_readiness_score']
                lbl=score_label(task_score)
                sc_hex='34c759' if task_score>=70 else ('ff9f0a' if task_score>=40 else 'ff3b30')
                sc_bg='e8f9ed' if task_score>=70 else ('fff4e0' if task_score>=40 else 'ffe5e3')

                tbl=doc.add_table(rows=1,cols=3); tbl.style='Table Grid'
                c0=tbl.rows[0].cells[0]; set_cell_bg(c0,sc_bg)
                p0=c0.paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_run(p0,f'{idx:02d}',bold=True,size=14,color='0071e3')
                c1=tbl.rows[0].cells[1]; set_cell_bg(c1,sc_bg)
                add_run(c1.paragraphs[0],task['name'],bold=True,size=11,color='1d1d1f')
                c2=tbl.rows[0].cells[2]; set_cell_bg(c2,sc_bg)
                p2=c2.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
                add_run(p2,f'{task_score:.0f}%  {lbl}',bold=True,size=11,color=sc_hex)
                set_col_width(c0,1.0); set_col_width(c1,10.5); set_col_width(c2,3.0)

                hrs_sv=result.get('estimated_hours_saved',0)
                dt=doc.add_table(rows=2,cols=2); dt.style='Table Grid'
                for ri,(k,v) in enumerate([('Difficulty',result.get('difficulty','N/A').title()),
                                            ('Hours Saved/Year',f'{hrs_sv:.1f} hrs')]):
                    kc=dt.rows[ri].cells[0]; vc=dt.rows[ri].cells[1]
                    set_cell_bg(kc,'f5f5f7'); bg='ffffff' if ri%2==0 else 'f5f5f7'; set_cell_bg(vc,bg)
                    set_col_width(kc,3.5); set_col_width(vc,11.0)
                    add_run(kc.paragraphs[0],k,bold=True,size=9,color='6e6e73')
                    add_run(vc.paragraphs[0],v,size=10,color='1d1d1f')
                if result.get('recommendation'):
                    doc.add_paragraph()
                    rt=doc.add_table(rows=1,cols=2); rt.style='Table Grid'
                    lc=rt.rows[0].cells[0]; rc=rt.rows[0].cells[1]
                    set_cell_bg(lc,'e8f1fc'); set_cell_bg(rc,'e8f1fc')
                    set_col_width(lc,2.5); set_col_width(rc,12.0)
                    add_run(lc.paragraphs[0],'Recommendation',bold=True,size=9,color='0071e3')
                    add_run(rc.paragraphs[0],result['recommendation'],size=10,color='1d1d1f')
                doc.add_paragraph()

        # Combined conclusion
        add_heading('Combined Summary')
        p=doc.add_paragraph()
        add_run(p,f'This combined report covers {len(analyses_list)} workflows with {total_tasks} total tasks. ',size=10,color='1d1d1f')
        add_run(p,f'Implementing automation recommendations could save {total_hours:.0f} hours annually',bold=True,size=10,color='1d1d1f')
        add_run(p,f', worth approximately ',size=10,color='1d1d1f')
        add_run(p,f'\u20ac{total_savings:,.0f}.',bold=True,size=10,color='1d1d1f')

        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        add_run(p,'Generated by WorkScanAI — AI-Powered Workflow Analysis',size=8,color='6e6e73')
        doc.save(output_path)
        return output_path

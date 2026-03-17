"""
Report generation service — full parity with website sections per context
Individual: Countdown Clock, Human Edge, Career Pivot
Team: Team Velocity, FTE, Rollout Timeline, 90-Day Sprint
Company: Competitor Gap, Headcount Signal, Industry Benchmark, Board Summary
All: F1 sub-scores, F3 risk flags, F4 readiness, F9 agentification
"""
import json as _json
import re as _re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, PageBreak, HRFlowable, KeepTogether)
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas as rl_canvas
from datetime import datetime
from typing import Dict, List

BLUE        = colors.HexColor('#0071e3')
BLUE_LIGHT  = colors.HexColor('#e8f1fc')
GRAY_900    = colors.HexColor('#1d1d1f')
GRAY_600    = colors.HexColor('#6e6e73')
GRAY_200    = colors.HexColor('#d2d2d7')
GRAY_100    = colors.HexColor('#f5f5f7')
WHITE       = colors.white
GREEN       = colors.HexColor('#34c759')
GREEN_LIGHT = colors.HexColor('#e8f9ed')
AMBER       = colors.HexColor('#ff9f0a')
AMBER_LIGHT = colors.HexColor('#fff4e0')
RED         = colors.HexColor('#ff3b30')
RED_LIGHT   = colors.HexColor('#ffe5e3')
PURPLE      = colors.HexColor('#5856d6')
PURPLE_LIGHT= colors.HexColor('#f0f0ff')
TEAL        = colors.HexColor('#30b0c7')
TEAL_LIGHT  = colors.HexColor('#e5f8fb')
ORANGE      = colors.HexColor('#ff6b35')
ORANGE_LIGHT= colors.HexColor('#fff0eb')


def score_color(score):
    if score >= 70: return GREEN, GREEN_LIGHT
    elif score >= 40: return AMBER, AMBER_LIGHT
    else: return RED, RED_LIGHT


def score_label(score):
    if score >= 70: return 'HIGH'
    elif score >= 40: return 'MEDIUM'
    else: return 'LOW'


def split_rec(text):
    """Split recommendation into (opt1, opt2) or (None, full_text)."""
    m1 = _re.search(r'(Option\s+1\s*[—\-–])', text)
    m2 = _re.search(r'(Option\s+2\s*[—\-–])', text)
    if m1 and m2:
        before = text[:m1.start()].strip()
        part1  = text[m1.start():m2.start()].strip()
        part2  = text[m2.start():].strip()
        combined = (f'{before}<br/><br/>' if before else '') + f'{part1}<br/><br/>{part2}'
        return combined, None
    elif m2:
        return f'{text[:m2.start()].strip()}<br/><br/>{text[m2.start():].strip()}', None
    return text, None


class NumberedCanvas(rl_canvas.Canvas):
    def __init__(self, *args, **kwargs):
        rl_canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_footer(num_pages)
            rl_canvas.Canvas.showPage(self)
        rl_canvas.Canvas.save(self)

    def draw_page_footer(self, page_count):
        self.saveState()
        w, h = A4
        self.setStrokeColor(GRAY_200); self.setLineWidth(0.5)
        self.line(18*mm, 16*mm, w - 18*mm, 16*mm)
        self.setFont('Helvetica', 8); self.setFillColor(GRAY_600)
        self.drawString(18*mm, 11*mm, 'WorkScanAI — AI-Powered Workflow Analysis')
        self.drawRightString(w - 18*mm, 11*mm, f'Page {self._pageNumber} of {page_count}')
        self.restoreState()


class ReportGenerator:

    # ─────────────────────────────────────────────────────────────────────────
    # PDF helpers shared across PDF methods
    # ─────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _pdf_task_blocks(sorted_results, analysis_data, W, s, ST, style_fn):
        """Return a list of KeepTogether blocks for detailed task analysis."""
        blocks = []
        hourly = analysis_data.get('hourly_rate', 50)
        for idx, result in enumerate(sorted_results, 1):
            task = result['task']
            task_score = result['ai_readiness_score']
            tc, tc_light = score_color(task_score)
            lbl = score_label(task_score)
            block = []

            # Header
            block.append(Table([[
                Paragraph(f'<font color="#0071e3"><b>{idx}</b></font>',
                    style_fn(f'tn{idx}', fontSize=22, fontName='Helvetica-Bold', textColor=BLUE, alignment=TA_CENTER)),
                Paragraph(f'<b>{task["name"]}</b>',
                    style_fn(f'th{idx}', fontSize=13, fontName='Helvetica-Bold', textColor=GRAY_900, leading=17)),
                Paragraph(f'<b>{task_score:.0f}%</b><br/><font size="8">{lbl}</font>',
                    style_fn(f'tb{idx}', fontSize=16, fontName='Helvetica-Bold', textColor=tc, alignment=TA_CENTER, leading=20)),
            ]], colWidths=[14*mm, W-14*mm-22*mm, 22*mm],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),tc_light),
                ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                ('LEFTPADDING',(0,0),(0,-1),10),('LEFTPADDING',(1,0),(1,-1),8),
                ('RIGHTPADDING',(2,0),(2,-1),10),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                ('LINEBELOW',(0,0),(-1,-1),1.5,tc)])))

            # Details rows
            hrs_sv = result.get('estimated_hours_saved', 0)
            val_sv = hrs_sv * hourly
            details = [
                ['Description',    task.get('description', '-')],
                ['Frequency',      task.get('frequency', 'N/A').capitalize()],
                ['Time per Task',  f'{task.get("time_per_task", 0)} min'],
                ['Category',       task.get('category', 'N/A').replace('_', ' ').title()],
                ['Implementation', result.get('difficulty', 'N/A').title()],
                ['Annual Savings', f'{hrs_sv:.1f} hrs  /  \u20ac{val_sv:,.0f}'],
            ]
            block.append(Table(
                [[Paragraph(f'<b>{k}</b>', ST['label']), Paragraph(v, ST['body'])] for k, v in details],
                colWidths=[42*mm, W-42*mm],
                style=TableStyle([('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('RIGHTPADDING',(0,0),(-1,-1),10),
                    ('VALIGN',(0,0),(-1,-1),'TOP'),
                    ('ROWBACKGROUNDS',(0,0),(-1,-1),[WHITE, GRAY_100]),
                    ('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))

            # Recommendation
            if result.get('recommendation'):
                rec_html, _ = split_rec(result['recommendation'])
                block.append(Table([[
                    Paragraph('<b>Recommendation</b>', style_fn(f'reclbl{idx}', fontSize=11,
                        fontName='Helvetica-Bold', textColor=BLUE, leading=14)),
                    Paragraph(rec_html, style_fn(f'rec{idx}', fontSize=11, fontName='Helvetica',
                        textColor=GRAY_900, leading=16, spaceAfter=0)),
                ]], colWidths=[42*mm, W-42*mm],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE_LIGHT),
                    ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('VALIGN',(0,0),(-1,-1),'TOP'),
                    ('LINEBEFORE',(0,0),(0,-1),3,BLUE)])))

            # F1 Sub-scores
            sub_keys = [('score_repeatability','Repeatability'),('score_data_availability','Data Avail.'),
                        ('score_error_tolerance','Error Tol.'),('score_integration','Integration')]
            sub_vals = [(lbl2, result.get(k)) for k, lbl2 in sub_keys if result.get(k) is not None]
            if sub_vals:
                sub_cells = [[
                    Paragraph(lbl2, style_fn(f'sl{idx}{lbl2}', fontSize=8, fontName='Helvetica', textColor=GRAY_600)),
                    Paragraph(f'<b>{val:.0f}</b>', style_fn(f'sv{idx}{lbl2}', fontSize=9,
                        fontName='Helvetica-Bold', textColor=GRAY_900, alignment=TA_RIGHT)),
                ] for lbl2, val in sub_vals]
                block.append(Table(sub_cells, colWidths=[W*0.55, W*0.45],
                    style=TableStyle([('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3),
                        ('LEFTPADDING',(0,0),(-1,-1),10),('RIGHTPADDING',(0,0),(-1,-1),10),
                        ('BACKGROUND',(0,0),(-1,-1),GRAY_100),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))

            # F3 Risk
            risk_flag = result.get('risk_flag')
            if risk_flag:
                risk_level = result.get('risk_level', 'safe')
                rc_map = {'safe':(GREEN,GREEN_LIGHT),'caution':(AMBER,AMBER_LIGHT),'warning':(RED,RED_LIGHT)}
                rc_fg, rc_bg = rc_map.get(risk_level, (GREEN, GREEN_LIGHT))
                block.append(Table([[
                    Paragraph('<b>Risk</b>', style_fn(f'rk{idx}', fontSize=9, fontName='Helvetica-Bold', textColor=rc_fg, leading=12)),
                    Paragraph(risk_flag, style_fn(f'rf{idx}', fontSize=9, fontName='Helvetica', textColor=GRAY_900, leading=13)),
                ]], colWidths=[18*mm, W-18*mm],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),rc_bg),
                    ('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                    ('LINEBEFORE',(0,0),(0,-1),3,rc_fg)])))

            # F9 Agentification
            agent_label = result.get('agent_label')
            if agent_label:
                phase_n = result.get('agent_phase', 1)
                ph_color = {1:AMBER,2:BLUE,3:GREEN}.get(phase_n, BLUE)
                agent_rows = [[
                    Paragraph('<b>Agentification</b>', style_fn(f'agl{idx}', fontSize=9, fontName='Helvetica-Bold', textColor=ph_color, leading=12)),
                    Paragraph(agent_label, style_fn(f'agv{idx}', fontSize=9, fontName='Helvetica-Bold', textColor=GRAY_900, leading=13)),
                ]]
                if result.get('agent_milestone'):
                    agent_rows.append([
                        Paragraph('Milestone', style_fn(f'aml{idx}', fontSize=8, fontName='Helvetica', textColor=GRAY_600, leading=12)),
                        Paragraph(result['agent_milestone'], style_fn(f'amv{idx}', fontSize=9, fontName='Helvetica', textColor=GRAY_900, leading=13)),
                    ])
                if result.get('orchestration'):
                    agent_rows.append([
                        Paragraph('Pipeline', style_fn(f'aol{idx}', fontSize=8, fontName='Helvetica', textColor=GRAY_600, leading=12)),
                        Paragraph(result['orchestration'], style_fn(f'aov{idx}', fontSize=9, fontName='Helvetica', textColor=GRAY_900, leading=13)),
                    ])
                block.append(Table(agent_rows, colWidths=[28*mm, W-28*mm],
                    style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE_LIGHT),
                        ('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
                        ('LEFTPADDING',(0,0),(-1,-1),10),('VALIGN',(0,0),(-1,-1),'TOP'),
                        ('LINEBEFORE',(0,0),(0,-1),3,ph_color),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))

            block.append(Spacer(1, 8*mm))
            blocks.append(KeepTogether(block))
        return blocks


    @staticmethod
    def _pdf_context_sections(story, analysis_data, W, style_fn, ST):
        """Append context-specific pages: Individual / Team / Company."""
        context = analysis_data.get('analysis_context', 'individual')
        results = analysis_data['results']
        sorted_results = sorted(results, key=lambda x: x['ai_readiness_score'], reverse=True)
        score = analysis_data['automation_score']
        hours = analysis_data['hours_saved']
        savings = analysis_data['annual_savings']
        avg_human_edge = sum(r.get('human_edge_score') or 50 for r in results) / max(len(results), 1)

        # ── INDIVIDUAL ────────────────────────────────────────────────────
        if context == 'individual':
            story.append(PageBreak())
            story.append(Paragraph('Career Future Analysis', ST['section_title']))
            story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=10))

            # B1 — Countdown Clock
            countdown_tasks = [r for r in sorted_results if r.get('countdown_window')]
            if countdown_tasks:
                story.append(Paragraph('Automation Countdown Clock', style_fn('cct', fontSize=13,
                    fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=6, spaceAfter=6)))
                cw_map = {
                    'now':   ('Automate NOW',   RED,  RED_LIGHT),
                    '12-24': ('12-24 months',   AMBER,AMBER_LIGHT),
                    '24-48': ('24-48 months',   colors.HexColor('#f5c518'), GRAY_100),
                    '48+':   ('48+ months',     GREEN,GREEN_LIGHT),
                }
                cw_rows = [[Paragraph('<b>Task</b>', ST['label']),
                             Paragraph('<b>Risk Window</b>', ST['label']),
                             Paragraph('<b>Human Edge</b>', ST['label'])]]
                for r in countdown_tasks:
                    cw = r.get('countdown_window', '24-48')
                    cw_label, cw_color, _ = cw_map.get(cw, ('Unknown', GRAY_600, GRAY_100))
                    he = r.get('human_edge_score')
                    cw_rows.append([
                        Paragraph(r['task']['name'], style_fn(f'cwn{r["task"]["name"][:8]}', fontSize=9, fontName='Helvetica', textColor=GRAY_900)),
                        Paragraph(cw_label, style_fn(f'cwl{cw}', fontSize=9, fontName='Helvetica-Bold', textColor=cw_color)),
                        Paragraph(f'{he:.0f}/100' if he else '—', style_fn(f'cwh{cw}', fontSize=9, fontName='Helvetica', textColor=GRAY_600)),
                    ])
                story.append(Table(cw_rows, colWidths=[W*0.45, W*0.35, W*0.20],
                    style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                        ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                        ('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6),
                        ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
                story.append(Spacer(1, 6*mm))
                story.append(Paragraph(
                    'The 900-Day Window: Within 900 days, any job done on a screen can be replaced by AI '
                    'for under \u20ac1,000/year. Tasks in the red zone are at immediate risk.',
                    style_fn('900d', fontSize=9, fontName='Helvetica-Oblique', textColor=GRAY_600, leading=14, spaceAfter=8)))

            # B2 — Human Edge
            story.append(Paragraph('Your Human Edge', style_fn('he_ttl', fontSize=13,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=10, spaceAfter=6)))
            he_rows = [
                [Paragraph('<b>Metric</b>', ST['label']), Paragraph('<b>Score</b>', ST['label'])],
                [Paragraph('AI Replacement Risk', style_fn('air', fontSize=9, fontName='Helvetica', textColor=GRAY_900)),
                 Paragraph(f'<font color="#ff3b30"><b>{score:.0f}%</b></font>', style_fn('airv', fontSize=11, fontName='Helvetica-Bold'))],
                [Paragraph('Human Irreplaceability', style_fn('hi', fontSize=9, fontName='Helvetica', textColor=GRAY_900)),
                 Paragraph(f'<font color="#ff9f0a"><b>{avg_human_edge:.0f}%</b></font>', style_fn('hiv', fontSize=11, fontName='Helvetica-Bold'))],
            ]
            story.append(Table(he_rows, colWidths=[W*0.6, W*0.4],
                style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                    ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                    ('TOPPADDING',(0,0),(-1,-1),7),('BOTTOMPADDING',(0,0),(-1,-1),7),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            insight = ('Your role has strong human-essential components. Amplify these while letting AI handle the rest.'
                       if avg_human_edge >= 60 else
                       'Your role is highly automatable. Now is the time to pivot to higher human-edge functions.')
            story.append(Paragraph(insight, style_fn('he_ins', fontSize=9, fontName='Helvetica-Oblique',
                textColor=GRAY_600, leading=14, spaceAfter=8, spaceBefore=4)))

            # B3 — Career Pivot
            pivot_tasks = [r for r in sorted_results if r.get('pivot_roles') or r.get('pivot_skills')]
            if pivot_tasks:
                story.append(Paragraph('Career Pivot Recommendations', style_fn('cpr', fontSize=13,
                    fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=10, spaceAfter=6)))

                # Aggregate skills (deduplicated)
                all_skills = []
                for r in pivot_tasks:
                    try:
                        sk = _json.loads(r['pivot_skills']) if isinstance(r.get('pivot_skills'), str) else (r.get('pivot_skills') or [])
                        for s in (sk if isinstance(sk, list) else []):
                            if s not in all_skills and len(all_skills) < 6:
                                all_skills.append(s)
                    except Exception: pass
                if all_skills:
                    story.append(Paragraph('Skills to Develop Now', style_fn('sk_ttl', fontSize=10,
                        fontName='Helvetica-Bold', textColor=BLUE, spaceAfter=4)))
                    story.append(Paragraph('  ·  '.join(all_skills), style_fn('sk_list', fontSize=9,
                        fontName='Helvetica', textColor=GRAY_900, leading=14, spaceAfter=6)))

                # Roles
                all_roles = []
                for r in pivot_tasks:
                    try:
                        roles = _json.loads(r['pivot_roles']) if isinstance(r.get('pivot_roles'), str) else (r.get('pivot_roles') or [])
                        for role in (roles if isinstance(roles, list) else []):
                            if not any(x.get('role') == role.get('role') for x in all_roles) and len(all_roles) < 5:
                                all_roles.append(role)
                    except Exception: pass
                if all_roles:
                    story.append(Paragraph('Adjacent Roles (Lower AI Risk)', style_fn('ro_ttl', fontSize=10,
                        fontName='Helvetica-Bold', textColor=BLUE, spaceAfter=4)))
                    role_rows = [[Paragraph('<b>Role</b>', ST['label']),
                                  Paragraph('<b>Risk</b>', ST['label']),
                                  Paragraph('<b>Pivot Distance</b>', ST['label'])]]
                    for role_item in all_roles:
                        rn = role_item.get('role', '—')
                        risk = role_item.get('risk', '—')
                        pdist = role_item.get('pivot_distance', '—')
                        rc2 = GREEN if risk == 'low' else (AMBER if risk == 'medium' else RED)
                        role_rows.append([
                            Paragraph(rn, style_fn(f'prn{rn[:8]}', fontSize=9, fontName='Helvetica', textColor=GRAY_900)),
                            Paragraph(risk, style_fn(f'prr{risk}', fontSize=9, fontName='Helvetica-Bold', textColor=rc2)),
                            Paragraph(pdist, style_fn(f'prd{pdist}', fontSize=9, fontName='Helvetica', textColor=GRAY_600)),
                        ])
                    story.append(Table(role_rows, colWidths=[W*0.5, W*0.2, W*0.3],
                        style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                            ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                            ('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
                            ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))

        # ── TEAM ──────────────────────────────────────────────────────────
        elif context == 'team':
            story.append(PageBreak())
            story.append(Paragraph('Team Automation Strategy', ST['section_title']))
            story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=10))

            # C1 — Velocity Impact
            story.append(Paragraph('Team Velocity Impact', style_fn('tv_ttl', fontSize=13,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=6, spaceAfter=6)))
            fte = hours / 1800
            vel_data = [
                ['Hours freed / yr', f'{hours:.0f}h', 'Available for product & growth'],
                ['FTE equivalent',   f'{fte:.1f}',    'Roles redeployable to strategic work'],
                ['Cost saved / yr',  f'\u20ac{savings:,.0f}', "At your team's hourly rate"],
            ]
            vel_rows = [[Paragraph('<b>Metric</b>', ST['label']),
                          Paragraph('<b>Value</b>', ST['label']),
                          Paragraph('<b>Note</b>', ST['label'])]]
            for metric, val, note in vel_data:
                vel_rows.append([
                    Paragraph(metric, style_fn(f'vm{metric[:6]}', fontSize=9, fontName='Helvetica-Bold', textColor=GRAY_900)),
                    Paragraph(f'<font color="#0071e3"><b>{val}</b></font>', style_fn(f'vv{metric[:6]}', fontSize=12, fontName='Helvetica-Bold')),
                    Paragraph(note, style_fn(f'vn{metric[:6]}', fontSize=9, fontName='Helvetica', textColor=GRAY_600)),
                ])
            story.append(Table(vel_rows, colWidths=[W*0.3, W*0.2, W*0.5],
                style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                    ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                    ('TOPPADDING',(0,0),(-1,-1),7),('BOTTOMPADDING',(0,0),(-1,-1),7),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            story.append(Spacer(1, 6*mm))

            # Rollout timeline
            story.append(Paragraph('Automation Rollout Timeline', style_fn('rt_ttl', fontSize=10,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceAfter=4)))
            phases_team = [
                ('Phase 1 — Quick Wins (0–3 months)',  [r for r in sorted_results if r.get('difficulty','').lower()=='easy'],   GREEN,  GREEN_LIGHT),
                ('Phase 2 — Medium-term (3–12 months)',[r for r in sorted_results if r.get('difficulty','').lower()=='medium'], AMBER,  AMBER_LIGHT),
                ('Phase 3 — Strategic (12–36 months)', [r for r in sorted_results if r.get('difficulty','').lower()=='hard'],   BLUE,   BLUE_LIGHT),
            ]
            for ph_name, ph_tasks, ph_col, ph_bg in phases_team:
                ph_hrs = sum(r.get('estimated_hours_saved',0) for r in ph_tasks)
                story.append(Table([[
                    Paragraph(f'<b>{ph_name}</b>', style_fn(f'pt{ph_name[:6]}', fontSize=10, fontName='Helvetica-Bold', textColor=ph_col)),
                    Paragraph(f'<b>{ph_hrs:.0f}h/yr</b>  ·  {len(ph_tasks)} tasks',
                        style_fn(f'pv{ph_name[:6]}', fontSize=10, fontName='Helvetica', textColor=GRAY_900, alignment=TA_RIGHT)),
                ]], colWidths=[W*0.65, W*0.35],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),ph_bg),
                    ('TOPPADDING',(0,0),(-1,-1),8),('BOTTOMPADDING',(0,0),(-1,-1),8),
                    ('LEFTPADDING',(0,0),(-1,-1),12),('RIGHTPADDING',(-1,0),(-1,-1),12),
                    ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEBELOW',(0,0),(-1,-1),1,ph_col)])))
            story.append(Spacer(1, 8*mm))

            # C2 — 90-Day Sprint Plan
            story.append(Paragraph('90-Day Sprint Plan', style_fn('sp_ttl', fontSize=13,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=10, spaceAfter=6)))
            sprint_tasks = sorted([r for r in results if r.get('difficulty','').lower()=='easy'],
                                   key=lambda x: x['ai_readiness_score'], reverse=True)[:5]
            if sprint_tasks:
                sp_rows = [[Paragraph('<b>#</b>', ST['label']),
                             Paragraph('<b>Task</b>', ST['label']),
                             Paragraph('<b>Score</b>', ST['label']),
                             Paragraph('<b>Hours/yr</b>', ST['label'])]]
                for i, r in enumerate(sprint_tasks, 1):
                    sp_rows.append([
                        Paragraph(str(i), style_fn(f'spn{i}', fontSize=9, fontName='Helvetica-Bold', textColor=GREEN)),
                        Paragraph(r['task']['name'], style_fn(f'spt{i}', fontSize=9, fontName='Helvetica', textColor=GRAY_900)),
                        Paragraph(f'{r["ai_readiness_score"]:.0f}%', style_fn(f'spsc{i}', fontSize=9, fontName='Helvetica-Bold', textColor=BLUE)),
                        Paragraph(f'{r.get("estimated_hours_saved",0):.0f}h', style_fn(f'sph{i}', fontSize=9, fontName='Helvetica', textColor=GRAY_900)),
                    ])
                story.append(Table(sp_rows, colWidths=[10*mm, W-10*mm-25*mm-30*mm, 25*mm, 30*mm],
                    style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                        ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                        ('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6),
                        ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            else:
                story.append(Paragraph('No easy-difficulty tasks — focus on Phase 2 medium-term automations.',
                    style_fn('sp_none', fontSize=9, fontName='Helvetica-Oblique', textColor=GRAY_600)))

        # ── COMPANY ───────────────────────────────────────────────────────
        elif context == 'company':
            story.append(PageBreak())
            story.append(Paragraph('Strategic Business Analysis', ST['section_title']))
            story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=10))

            # D1 — Competitor Gap
            story.append(Paragraph('AI-First Competitor Gap', style_fn('cg_ttl', fontSize=13,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=6, spaceAfter=6)))
            comp_data = [
                ('If you automate now',      savings,         'Your annual advantage',                    GREEN, GREEN_LIGHT),
                ('If you wait 12 months',    savings * 0.35,  '65% of advantage lost to delayed adoption',AMBER, AMBER_LIGHT),
                ('AI-first competitor edge', savings * 1.4,   'Over you if they move first',              RED,   RED_LIGHT),
            ]
            comp_rows = []
            comp_hex = {GREEN: '34c759', AMBER: 'ff9f0a', RED: 'ff3b30'}
            for label, val, note, c_fg, c_bg in comp_data:
                comp_rows.append([
                    Paragraph(f'<b>{label}</b>', style_fn(f'cl{label[:6]}', fontSize=9, fontName='Helvetica-Bold', textColor=GRAY_900)),
                    Paragraph(f'<font color="#{comp_hex.get(c_fg,"0071e3")}"><b>\u20ac{val:,.0f}/yr</b></font>',
                        style_fn(f'cv{label[:6]}', fontSize=12, fontName='Helvetica-Bold')),
                    Paragraph(note, style_fn(f'cn{label[:6]}', fontSize=9, fontName='Helvetica', textColor=GRAY_600)),
                ])
            story.append(Table(comp_rows, colWidths=[W*0.32, W*0.25, W*0.43],
                style=TableStyle([('ROWBACKGROUNDS',(0,0),(-1,-1),[GREEN_LIGHT,AMBER_LIGHT,RED_LIGHT]),
                    ('TOPPADDING',(0,0),(-1,-1),8),('BOTTOMPADDING',(0,0),(-1,-1),8),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            story.append(Spacer(1, 6*mm))

            # D2 — Headcount Signal
            story.append(Paragraph('Headcount Signal', style_fn('hc_ttl', fontSize=13,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=10, spaceAfter=6)))
            fte2 = hours / 1800
            hc_data = [
                ['Hours freed / yr',    f'{hours:.0f}h',         'Total across all tasks'],
                ['FTE equivalent',      f'{fte2:.1f}',           'At 1,800 working hrs/yr'],
                ['Saved per FTE',       f'\u20ac{savings/max(fte2,0.1):,.0f}', 'Annual cost per role'],
            ]
            hc_rows = [[Paragraph('<b>Metric</b>', ST['label']),
                         Paragraph('<b>Value</b>', ST['label']),
                         Paragraph('<b>Note</b>', ST['label'])]]
            for m, v, n in hc_data:
                hc_rows.append([
                    Paragraph(m, style_fn(f'hm{m[:6]}', fontSize=9, fontName='Helvetica-Bold', textColor=GRAY_900)),
                    Paragraph(f'<font color="#5856d6"><b>{v}</b></font>', style_fn(f'hv{m[:6]}', fontSize=12, fontName='Helvetica-Bold')),
                    Paragraph(n, style_fn(f'hn{m[:6]}', fontSize=9, fontName='Helvetica', textColor=GRAY_600)),
                ])
            story.append(Table(hc_rows, colWidths=[W*0.3, W*0.2, W*0.5],
                style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                    ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                    ('TOPPADDING',(0,0),(-1,-1),7),('BOTTOMPADDING',(0,0),(-1,-1),7),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            story.append(Paragraph('Recommended: redeploy freed capacity to AI oversight, customer relationships, '
                'and strategic growth — not headcount reduction.',
                style_fn('hc_note', fontSize=9, fontName='Helvetica-Oblique', textColor=GRAY_600, leading=13,
                    spaceAfter=8, spaceBefore=4)))

            # D3 — Industry Benchmark
            story.append(Paragraph('Industry Benchmark', style_fn('ib_ttl', fontSize=13,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=10, spaceAfter=6)))
            gap = max(0, 81 - round(score))
            bm_data = [
                ('Your score',        f'{score:.0f}%', 'This workflow',         BLUE),
                ('Sector average',    '58%',           'Cognitive workflows',   GRAY_600),
                ('Top 10% orgs',      '81%',           'AI-first companies',    GREEN),
                ('Gap to close',      f'{gap}%',       'To reach top 10%',      AMBER),
            ]
            bm_rows = [[Paragraph('<b>Benchmark</b>', ST['label']),
                         Paragraph('<b>Score</b>', ST['label']),
                         Paragraph('<b>Context</b>', ST['label'])]]
            for bm_label, bm_val, bm_note, bm_col in bm_data:
                bm_rows.append([
                    Paragraph(bm_label, style_fn(f'bml{bm_label[:6]}', fontSize=9, fontName='Helvetica-Bold', textColor=GRAY_900)),
                    Paragraph(f'<b>{bm_val}</b>', style_fn(f'bmv{bm_label[:6]}', fontSize=12, fontName='Helvetica-Bold', textColor=bm_col)),
                    Paragraph(bm_note, style_fn(f'bmn{bm_label[:6]}', fontSize=9, fontName='Helvetica', textColor=GRAY_600)),
                ])
            story.append(Table(bm_rows, colWidths=[W*0.35, W*0.2, W*0.45],
                style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                    ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                    ('TOPPADDING',(0,0),(-1,-1),7),('BOTTOMPADDING',(0,0),(-1,-1),7),
                    ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            story.append(Spacer(1, 6*mm))

            # D4 — Board Summary
            story.append(Paragraph('Board-Ready Executive Summary', style_fn('bs_ttl', fontSize=13,
                fontName='Helvetica-Bold', textColor=GRAY_900, spaceBefore=10, spaceAfter=6)))
            quick_wins = len([r for r in results if r.get('difficulty','').lower()=='easy'])
            risk_flags  = len([r for r in results if r.get('risk_level') != 'safe'])
            wf_name = analysis_data['workflow'].get('name', 'Workflow')
            wf_ind  = analysis_data['workflow'].get('industry', '')
            board_lines = [
                ('Workflow',            wf_name),
                ('Industry',            wf_ind or 'General'),
                ('Automation potential',f'{score:.0f}% of workflow tasks'),
                ('Annual savings',      f'\u20ac{savings:,.0f}'),
                ('Hours reclaimed',     f'{hours:.0f}h/yr'),
                ('FTE equivalent',      f'{hours/1800:.1f} roles'),
                ('Quick wins (90 days)',f'{quick_wins} tasks'),
                ('Risk flags',          f'{risk_flags} tasks require compliance review'),
                ('Recommendation',      'Begin 90-day automation sprint. Prioritise quick wins. Redeploy freed capacity to strategic functions.'),
            ]
            board_rows = []
            for k, v in board_lines:
                board_rows.append([
                    Paragraph(f'<b>{k}</b>', style_fn(f'bk{k[:6]}', fontSize=9, fontName='Helvetica-Bold', textColor=GRAY_600)),
                    Paragraph(v, style_fn(f'bv{k[:6]}', fontSize=9, fontName='Helvetica', textColor=GRAY_900)),
                ])
            story.append(Table(board_rows, colWidths=[W*0.38, W*0.62],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),colors.HexColor('#1d1d1f')),
                    ('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6),
                    ('LEFTPADDING',(0,0),(-1,-1),12),('LINEBELOW',(0,0),(-1,-1),0.3,colors.HexColor('#3a3a3c')),
                    ('TEXTCOLOR',(0,0),(0,-1),colors.HexColor('#86868b')),
                    ('TEXTCOLOR',(1,0),(1,-1),colors.white),
                ])))

    # ─────────────────────────────────────────────────────────────────────────
    # MAIN PDF GENERATOR
    # ─────────────────────────────────────────────────────────────────────────

    @staticmethod
    def generate_pdf_report(analysis_data: Dict, output_path: str):
        doc = SimpleDocTemplate(output_path, pagesize=A4,
            leftMargin=18*mm, rightMargin=18*mm, topMargin=20*mm, bottomMargin=24*mm)
        W = A4[0] - 36*mm
        story = []
        s = getSampleStyleSheet()

        def style(name, parent='Normal', **kw):
            return ParagraphStyle(name, parent=s[parent], **kw)

        ST = {
            'section_title': style('st', fontSize=18, leading=22, textColor=GRAY_900,
                                    spaceBefore=18, spaceAfter=8, fontName='Helvetica-Bold'),
            'label':  style('lb', fontSize=9, leading=12, textColor=GRAY_600,
                             fontName='Helvetica-Bold', spaceAfter=2),
            'body':   style('bo', fontSize=10, leading=15, textColor=GRAY_900,
                             fontName='Helvetica', spaceAfter=6),
            'cover_sub':  style('cs', fontSize=13, leading=18, textColor=GRAY_600,
                                 spaceAfter=4, fontName='Helvetica'),
            'cover_meta': style('cm', fontSize=10, leading=14, textColor=GRAY_600,
                                 fontName='Helvetica'),
        }

        workflow = analysis_data['workflow']
        score    = analysis_data['automation_score']
        hours    = analysis_data['hours_saved']
        savings  = analysis_data['annual_savings']
        results  = analysis_data['results']
        sorted_results = sorted(results, key=lambda x: x['ai_readiness_score'], reverse=True)
        high = [r for r in results if r['ai_readiness_score'] >= 70]
        med  = [r for r in results if 40 <= r['ai_readiness_score'] < 70]
        low  = [r for r in results if r['ai_readiness_score'] < 40]
        sc, sc_light = score_color(score)

        # ── Cover ─────────────────────────────────────────────────────────
        story.append(Table([['']],colWidths=[W],rowHeights=[4],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE)])))
        story.append(Spacer(1, 16*mm))
        story.append(Paragraph('WorkScanAI', style('brand', fontSize=11, textColor=BLUE,
            fontName='Helvetica-Bold', spaceAfter=10)))

        context = analysis_data.get('analysis_context', 'individual')
        ctx_label = {'individual': 'Personal Career Analysis', 'team': 'Team / Startup Analysis',
                     'company': 'Company / Department Analysis'}.get(context, 'Workflow Analysis')
        story.append(Paragraph(ctx_label, style('ctx_lbl', fontSize=11, textColor=GRAY_600,
            fontName='Helvetica', spaceAfter=6)))
        story.append(Paragraph('Workflow Automation Analysis Report',
            style('ct', fontSize=22, leading=27, textColor=GRAY_900, spaceAfter=6, fontName='Helvetica-Bold')))
        story.append(Paragraph(workflow['name'], style('wn', fontSize=22, leading=26,
            textColor=BLUE, fontName='Helvetica-Bold', spaceAfter=10)))
        if workflow.get('description'):
            story.append(Paragraph(workflow['description'], ST['cover_sub']))
        if workflow.get('industry'):
            story.append(Paragraph(f"Industry: {workflow['industry']}",
                style('ind', fontSize=9, textColor=GRAY_600, fontName='Helvetica', spaceAfter=4)))

        source_text = workflow.get('source_text', '').strip()
        if source_text:
            mode_labels = {'voice':'Voice Transcript','document':'Extracted Document Text','manual':'Original Input'}
            story.append(Spacer(1, 4*mm))
            story.append(Paragraph(mode_labels.get(workflow.get('input_mode','manual'),'Original Input'),
                style('src_lbl', fontSize=9, leading=12, textColor=BLUE, fontName='Helvetica-Bold', spaceAfter=4)))
            story.append(Table([[Paragraph(source_text,
                style('src', fontSize=9, leading=14, textColor=GRAY_900, fontName='Helvetica'))]],
                colWidths=[W], style=TableStyle([('BACKGROUND',(0,0),(-1,-1),GRAY_100),
                    ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                    ('LEFTPADDING',(0,0),(-1,-1),14),('RIGHTPADDING',(0,0),(-1,-1),14),
                    ('LINEBEFORE',(0,0),(0,-1),3,BLUE)])))

        story.append(Spacer(1, 8*mm))
        story.append(Paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')}", ST['cover_meta']))
        story.append(Spacer(1, 14*mm))

        hero = Table([[
            Paragraph(f'<font size="48"><b>{score:.0f}%</b></font><br/><font size="14" color="#6e6e73">Automation Score</font>',
                       style('hs', fontName='Helvetica-Bold', alignment=TA_CENTER, leading=54)),
            Table([[Paragraph('<b>Annual Savings</b>', ST['label']),''],
                   [Paragraph(f'\u20ac{savings:,.0f}', style('sv', fontSize=22, fontName='Helvetica-Bold', textColor=GRAY_900)),''],
                   [Paragraph(f'{hours:.0f} hours reclaimed', style('hr', fontSize=10, textColor=GRAY_600, fontName='Helvetica')),'']],
                  colWidths=[W*0.35, W*0.05]),
            Table([[Paragraph('<b>Tasks Analyzed</b>', ST['label']),''],
                   [Paragraph(str(len(results)), style('tc', fontSize=22, fontName='Helvetica-Bold', textColor=GRAY_900)),''],
                   [Paragraph('workflow tasks', style('tl', fontSize=10, textColor=GRAY_600, fontName='Helvetica')),'']],
                  colWidths=[W*0.25, W*0.05]),
        ]], colWidths=[W*0.28, W*0.42, W*0.30],
        style=TableStyle([('BACKGROUND',(0,0),(-1,-1),sc_light),
            ('TOPPADDING',(0,0),(-1,-1),16),('BOTTOMPADDING',(0,0),(-1,-1),16),
            ('LEFTPADDING',(0,0),(0,-1),20),('LEFTPADDING',(1,0),(1,-1),10),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEAFTER',(0,0),(0,-1),1,sc)]))
        story.append(hero)
        story.append(Spacer(1, 10*mm))
        story.append(Table([[
            Paragraph(f'<font color="#34c759"><b>{len(high)} HIGH</b></font> potential', ST['body']),
            Paragraph(f'<font color="#ff9f0a"><b>{len(med)} MEDIUM</b></font> potential', ST['body']),
            Paragraph(f'<font color="#ff3b30"><b>{len(low)} LOW</b></font> potential', ST['body']),
        ]], colWidths=[W/3,W/3,W/3],
        style=TableStyle([('BACKGROUND',(0,0),(-1,-1),GRAY_100),
            ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
            ('LEFTPADDING',(0,0),(-1,-1),14),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('LINEAFTER',(0,0),(1,-1),0.5,GRAY_200)])))
        story.append(PageBreak())

        # ── F4 AI Readiness ───────────────────────────────────────────────
        rs = analysis_data.get('readiness_score')
        rd_labels = [('Data Quality',analysis_data.get('readiness_data_quality')),
                     ('Process Clarity',analysis_data.get('readiness_process_docs')),
                     ('Tool Maturity',analysis_data.get('readiness_tool_maturity')),
                     ('Error Tolerance',analysis_data.get('readiness_team_skills'))]
        rd_vals = [(l,v) for l,v in rd_labels if v is not None]
        if rs is not None or rd_vals:
            story.append(Paragraph('AI Readiness Assessment', ST['section_title']))
            story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=8))
            if rs is not None:
                rs_color, rs_bg = score_color(rs)
                story.append(Table([[
                    Paragraph(f'<font size="32"><b>{rs:.0f}</b></font><font size="12" color="#6e6e73"> / 100</font><br/>'
                               f'<font size="10" color="#6e6e73">Overall AI Readiness</font>',
                        style('rso', fontName='Helvetica-Bold', alignment=TA_CENTER, leading=38)),
                    Paragraph('How prepared your organisation is for AI automation based on data quality, '
                               'process documentation, tool maturity and team skills.',
                        style('rsd', fontSize=9, fontName='Helvetica', textColor=GRAY_600, leading=14)),
                ]], colWidths=[W*0.3, W*0.7],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),rs_bg),
                    ('TOPPADDING',(0,0),(-1,-1),14),('BOTTOMPADDING',(0,0),(-1,-1),14),
                    ('LEFTPADDING',(0,0),(-1,-1),14),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                    ('LINEAFTER',(0,0),(0,-1),1,rs_color)])))
                story.append(Spacer(1,6*mm))
            if rd_vals:
                rd_rows = []
                for lbl2, val in rd_vals:
                    rc2, _ = score_color(val)
                    rd_rows.append([Paragraph(lbl2, style(f'rdl{lbl2}', fontSize=9,
                        fontName='Helvetica-Bold', textColor=GRAY_600)),
                        Paragraph(f'<b>{val:.0f}</b>', style(f'rdv{lbl2}', fontSize=10,
                        fontName='Helvetica-Bold', textColor=rc2))])
                story.append(Table(rd_rows, colWidths=[W*0.5, W*0.5],
                    style=TableStyle([('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6),
                        ('LEFTPADDING',(0,0),(-1,-1),14),
                        ('ROWBACKGROUNDS',(0,0),(-1,-1),[WHITE,GRAY_100]),
                        ('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
                story.append(Spacer(1, 10*mm))

        # ── Detailed Task Analysis ────────────────────────────────────────
        story.append(Paragraph('Detailed Task Analysis', ST['section_title']))
        story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=10))
        for blk in ReportGenerator._pdf_task_blocks(sorted_results, analysis_data, W, s, ST, style):
            story.append(blk)

        # ── Implementation Roadmap ────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Implementation Roadmap', ST['section_title']))
        story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=10))
        phases = [
            ('Phase 1 — Quick Wins','0–3 months','High score + easy setup = immediate ROI.',
             GREEN,GREEN_LIGHT,[r for r in sorted_results if r['ai_readiness_score']>=70 and r.get('difficulty','').lower()=='easy']),
            ('Phase 2 — Medium-Term','3–6 months','Technical setup required but significant savings.',
             AMBER,AMBER_LIGHT,[r for r in sorted_results if r['ai_readiness_score']>=50 and r.get('difficulty','').lower()=='medium']),
            ('Phase 3 — Advanced','6–12 months','Complex automations and custom development.',
             BLUE,BLUE_LIGHT,[r for r in sorted_results if r['ai_readiness_score']>=40 and r.get('difficulty','').lower()=='hard']),
        ]
        for title,timeline,desc,col,col_light,phase_tasks in phases:
            pb = []
            pb.append(Table([[
                Paragraph(f'<b>{title}</b>', style(f'ph{title}', fontSize=13,
                    fontName='Helvetica-Bold', textColor=col)),
                Paragraph(timeline, style(f'pt{title}', fontSize=10, textColor=GRAY_600,
                    fontName='Helvetica', alignment=TA_RIGHT)),
            ]], colWidths=[W*0.65, W*0.35],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),col_light),
                ('TOPPADDING',(0,0),(-1,-1),10),('BOTTOMPADDING',(0,0),(-1,-1),10),
                ('LEFTPADDING',(0,0),(0,-1),12),('RIGHTPADDING',(-1,0),(-1,-1),12),
                ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEBELOW',(0,0),(-1,-1),1.5,col)])))
            pb.append(Paragraph(desc, style(f'pd{title}', fontSize=9, textColor=GRAY_600,
                fontName='Helvetica', spaceAfter=6, leftIndent=12)))
            if phase_tasks:
                for r in phase_tasks:
                    pb.append(Table([[
                        Paragraph('>', style(f'ar{r["task"]["name"][:8]}', fontSize=10,
                            fontName='Helvetica-Bold', textColor=col)),
                        Paragraph(f'<b>{r["task"]["name"]}</b>  <font size="9" color="#6e6e73">'
                            f'Score {r["ai_readiness_score"]:.0f}%  {r.get("estimated_hours_saved",0):.0f} hrs saved</font>',
                            style(f'ri{r["task"]["name"][:8]}', fontSize=10, fontName='Helvetica',
                                textColor=GRAY_900, leading=14)),
                    ]], colWidths=[8*mm, W-8*mm],
                    style=TableStyle([('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
                        ('LEFTPADDING',(0,0),(-1,-1),12),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
            else:
                pb.append(Paragraph('No tasks in this phase.', style(f'np{title}', fontSize=9,
                    textColor=GRAY_600, fontName='Helvetica-Oblique', leftIndent=12, spaceAfter=4)))
            pb.append(Spacer(1, 8*mm))
            story.append(KeepTogether(pb))

        # ── Context-specific sections ─────────────────────────────────────
        ReportGenerator._pdf_context_sections(story, analysis_data, W, style, ST)

        # ── Conclusion ────────────────────────────────────────────────────
        story.append(PageBreak())
        story.append(Paragraph('Conclusion', ST['section_title']))
        story.append(HRFlowable(width=W, thickness=0.5, color=GRAY_200, spaceAfter=10))
        story.append(Paragraph(
            f'This analysis identified automation opportunities across <b>{len(results)} tasks</b> '
            f'in <b>{workflow["name"]}</b>. Implementing recommendations could save '
            f'<b>{hours:.0f} hours annually</b>, worth approximately <b>\u20ac{savings:,.0f}</b>.',
            ST['body']))
        story.append(Spacer(1, 6*mm))
        for item in ['Begin with Phase 1 quick wins — ROI within weeks, not months.',
                     'Use human-in-the-loop for tasks scored 40–70 to reduce risk.',
                     'Track time saved monthly to measure and communicate impact.',
                     'Revisit quarterly — new AI tools appear constantly.',
                     'Train team members on newly automated workflows.']:
            story.append(Table([[
                Paragraph('>', style(f'ai{item[:10]}', fontSize=11, textColor=BLUE, fontName='Helvetica-Bold')),
                Paragraph(item, ST['body']),
            ]], colWidths=[8*mm, W-8*mm],
            style=TableStyle([('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
                ('LEFTPADDING',(0,0),(0,-1),0),('VALIGN',(0,0),(-1,-1),'TOP')])))

        doc.build(story, canvasmaker=NumberedCanvas)
        return output_path


    # ─────────────────────────────────────────────────────────────────────────
    # DOCX GENERATOR
    # ─────────────────────────────────────────────────────────────────────────

    @staticmethod
    def generate_docx_report(analysis_data: Dict, output_path: str):
        if Document is None:
            raise ImportError("python-docx not installed")
        doc = Document()
        for sec in doc.sections:
            sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
            sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

        workflow = analysis_data['workflow']
        score    = analysis_data['automation_score']
        hours    = analysis_data['hours_saved']
        savings  = analysis_data['annual_savings']
        results  = analysis_data['results']
        context  = analysis_data.get('analysis_context', 'individual')
        sorted_results = sorted(results, key=lambda x: x['ai_readiness_score'], reverse=True)
        hourly   = analysis_data.get('hourly_rate', 50)

        def set_bg(cell, hex_color):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
            tcPr.append(shd)

        def set_w(cell, width_cm):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width_cm*567))); tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        def run(para, text, bold=False, size=None, color=None, italic=False):
            r = para.add_run(text); r.bold=bold; r.italic=italic
            if size: r.font.size = Pt(size)
            if color: r.font.color.rgb = RGBColor(*tuple(int(color[i:i+2],16) for i in (0,2,4)))
            return r

        def heading(text, size=16, color='1d1d1f'):
            p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
            r2 = p.add_run(text); r2.bold=True; r2.font.size=Pt(size)
            r2.font.color.rgb=RGBColor(*tuple(int(color[i:i+2],16) for i in (0,2,4)))
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
            bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'4')
            bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'d2d2d7')
            pBdr.append(bottom); pPr.append(pBdr)

        def kv_table(rows_data, col_widths=(4.0,11.0)):
            """rows_data: list of (key_str, val_str, bg_hex_or_None)"""
            t = doc.add_table(rows=len(rows_data), cols=2); t.style='Table Grid'
            for ri, (k, v, bg) in enumerate(rows_data):
                kc=t.rows[ri].cells[0]; vc=t.rows[ri].cells[1]
                set_bg(kc,'f5f5f7'); set_bg(vc, bg or ('ffffff' if ri%2==0 else 'f5f5f7'))
                set_w(kc, col_widths[0]); set_w(vc, col_widths[1])
                run(kc.paragraphs[0], k, bold=True, size=9, color='6e6e73')
                run(vc.paragraphs[0], v, size=10, color='1d1d1f')
            return t

        # ── Cover ─────────────────────────────────────────────────────────
        p = doc.add_paragraph(); run(p, 'WorkScanAI', bold=True, size=11, color='0071e3')
        ctx_label = {'individual':'Personal Career Analysis','team':'Team / Startup Analysis',
                     'company':'Company / Department Analysis'}.get(context,'Workflow Analysis')
        p = doc.add_paragraph(); run(p, ctx_label, size=10, color='86868b')
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
        p = doc.add_paragraph(); run(p,'Workflow Automation Analysis Report',bold=True,size=22,color='1d1d1f')
        p = doc.add_paragraph(); run(p,workflow['name'],bold=True,size=20,color='0071e3')
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
        if workflow.get('description'):
            p=doc.add_paragraph(); run(p,workflow['description'],size=11,color='6e6e73')
        if workflow.get('industry'):
            p=doc.add_paragraph(); run(p,f"Industry: {workflow['industry']}",size=9,color='86868b')

        source_text = workflow.get('source_text','').strip()
        if source_text:
            mode_labels={'voice':'Voice Transcript','document':'Extracted Document Text','manual':'Original Input'}
            p=doc.add_paragraph(); run(p,mode_labels.get(workflow.get('input_mode','manual'),'Input'),bold=True,size=9,color='0071e3')
            st=doc.add_table(rows=1,cols=1); st.style='Table Grid'
            cell=st.rows[0].cells[0]; set_bg(cell,'f5f5f7')
            run(cell.paragraphs[0],source_text,size=10,color='1d1d1f')
            doc.add_paragraph()

        p=doc.add_paragraph(); run(p,f"Generated {datetime.now().strftime('%B %d, %Y')}",size=9,color='6e6e73')
        p.paragraph_format.space_after=Pt(12)

        # Hero stats table
        t=doc.add_table(rows=2,cols=4); t.style='Table Grid'
        for i,(h,v,bg) in enumerate([('Automation Score',f'{score:.0f}%','e8f1fc'),
                                       ('Annual Savings',f'\u20ac{savings:,.0f}','e8f9ed'),
                                       ('Hours Reclaimed',f'{hours:.0f} hrs','fff4e0'),
                                       ('Tasks Analyzed',str(len(results)),'f5f5f7')]):
            hc=t.rows[0].cells[i]; vc=t.rows[1].cells[i]
            set_bg(hc,bg); set_bg(vc,bg)
            ph=hc.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(ph,h,bold=True,size=8,color='6e6e73')
            pv=vc.paragraphs[0]; pv.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(pv,v,bold=True,size=16,color='1d1d1f')
        doc.add_paragraph()


        # ── F4 AI Readiness ───────────────────────────────────────────────
        rs = analysis_data.get('readiness_score')
        if rs is not None:
            heading('AI Readiness Assessment', size=14, color='0071e3')
            rd_rows = [
                ('Overall Readiness',  f'{rs:.0f} / 100',                                    None),
                ('Data Quality',       f'{analysis_data.get("readiness_data_quality") or "—":.0f}' if analysis_data.get("readiness_data_quality") else '—', None),
                ('Process Clarity',    f'{analysis_data.get("readiness_process_docs") or "—":.0f}' if analysis_data.get("readiness_process_docs") else '—', None),
                ('Tool Maturity',      f'{analysis_data.get("readiness_tool_maturity") or "—":.0f}' if analysis_data.get("readiness_tool_maturity") else '—', None),
                ('Error Tolerance',    f'{analysis_data.get("readiness_team_skills") or "—":.0f}'   if analysis_data.get("readiness_team_skills")  else '—', None),
            ]
            kv_table(rd_rows)
            doc.add_paragraph()

        # ── Task Breakdown ────────────────────────────────────────────────
        heading('Task Breakdown')
        for idx, result in enumerate(sorted_results, 1):
            task = result['task']; task_score = result['ai_readiness_score']
            lbl = score_label(task_score)
            s_hex = '34c759' if task_score>=70 else ('ff9f0a' if task_score>=40 else 'ff3b30')
            bg_hex = 'e8f9ed' if task_score>=70 else ('fff4e0' if task_score>=40 else 'ffe5e3')

            # Header row
            t=doc.add_table(rows=1,cols=3); t.style='Table Grid'
            c0=t.rows[0].cells[0]; set_bg(c0,bg_hex); p2=c0.paragraphs[0]
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; run(p2,f'{idx:02d}',bold=True,size=18,color='0071e3')
            c1=t.rows[0].cells[1]; set_bg(c1,bg_hex); run(c1.paragraphs[0],task['name'],bold=True,size=13,color='1d1d1f')
            c2=t.rows[0].cells[2]; set_bg(c2,bg_hex); p3=c2.paragraphs[0]
            p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            run(p3,f'{task_score:.0f}%\n',bold=True,size=16,color=s_hex)
            run(p3,lbl,bold=True,size=8,color=s_hex)
            set_w(c0,1.2); set_w(c1,11.0); set_w(c2,2.8)
            doc.add_paragraph()

            # Details
            hrs_sv = result.get('estimated_hours_saved',0); val_sv=hrs_sv*hourly
            kv_table([
                ('Description',   task.get('description','-'), None),
                ('Frequency',     task.get('frequency','N/A').capitalize(), None),
                ('Time per Task', f'{task.get("time_per_task",0)} minutes', None),
                ('Category',      task.get('category','N/A').replace('_',' ').title(), None),
                ('Difficulty',    result.get('difficulty','N/A').title(), None),
                ('Annual Savings',f'{hrs_sv:.1f} hrs  /  \u20ac{val_sv:,.0f}', None),
            ])

            # F1 Sub-scores
            sub_keys = [('score_repeatability','Repeatability'),('score_data_availability','Data Availability'),
                        ('score_error_tolerance','Error Tolerance'),('score_integration','Integration Ease')]
            sub_vals = [(lbl2, result.get(k)) for k, lbl2 in sub_keys if result.get(k) is not None]
            if sub_vals:
                doc.add_paragraph()
                p=doc.add_paragraph(); run(p,'Sub-scores:',bold=True,size=9,color='86868b')
                t2=doc.add_table(rows=1,cols=len(sub_vals)); t2.style='Table Grid'
                for ci,(lbl2,val) in enumerate(sub_vals):
                    c=t2.rows[0].cells[ci]; set_bg(c,'f5f5f7')
                    ph=c.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    run(ph,f'{val:.0f}\n',bold=True,size=14,color=('34c759' if val>=70 else ('ff9f0a' if val>=45 else 'ff3b30')))
                    run(ph,lbl2,size=8,color='6e6e73')

            # Recommendation
            if result.get('recommendation'):
                doc.add_paragraph()
                rt=doc.add_table(rows=1,cols=2); rt.style='Table Grid'
                lc=rt.rows[0].cells[0]; rc=rt.rows[0].cells[1]
                set_bg(lc,'e8f1fc'); set_bg(rc,'e8f1fc'); set_w(lc,3.8); set_w(rc,11.2)
                run(lc.paragraphs[0],'Recommendation',bold=True,size=12,color='0071e3')
                rec_text = result['recommendation']
                m1=_re.search(r'(Option\s+1\s*[—\-–])',rec_text); m2=_re.search(r'(Option\s+2\s*[—\-–])',rec_text)
                if m1 and m2:
                    before=rec_text[:m1.start()].strip(); p1t=rec_text[m1.start():m2.start()].strip(); p2t=rec_text[m2.start():].strip()
                    if before: run(rc.paragraphs[0],before,size=11,color='1d1d1f'); pm=rc.add_paragraph(); run(pm,p1t,size=11,color='1d1d1f')
                    else: run(rc.paragraphs[0],p1t,size=11,color='1d1d1f')
                    run(rc.add_paragraph(),p2t,size=11,color='1d1d1f')
                elif m2:
                    run(rc.paragraphs[0],rec_text[:m2.start()].strip(),size=11,color='1d1d1f')
                    run(rc.add_paragraph(),rec_text[m2.start():].strip(),size=11,color='1d1d1f')
                else:
                    run(rc.paragraphs[0],rec_text,size=11,color='1d1d1f')

            # F3 Risk
            if result.get('risk_flag'):
                doc.add_paragraph()
                rt2=doc.add_table(rows=1,cols=2); rt2.style='Table Grid'
                rl_cell=rt2.rows[0].cells[0]; rv_cell=rt2.rows[0].cells[1]
                risk_lv = result.get('risk_level','safe')
                risk_bg = 'e8f9ed' if risk_lv=='safe' else ('fff4e0' if risk_lv=='caution' else 'ffe5e3')
                risk_col= '34c759' if risk_lv=='safe' else ('ff9f0a' if risk_lv=='caution' else 'ff3b30')
                set_bg(rl_cell,risk_bg); set_bg(rv_cell,risk_bg); set_w(rl_cell,2.0); set_w(rv_cell,13.0)
                run(rl_cell.paragraphs[0],'Risk',bold=True,size=9,color=risk_col)
                run(rv_cell.paragraphs[0],result['risk_flag'],size=9,color='1d1d1f')

            # F9 Agentification
            if result.get('agent_label'):
                doc.add_paragraph()
                rt3=doc.add_table(rows=1,cols=2); rt3.style='Table Grid'
                al=rt3.rows[0].cells[0]; av=rt3.rows[0].cells[1]
                ph_col = '34c759' if result.get('agent_phase')==3 else ('0071e3' if result.get('agent_phase')==2 else 'ff9f0a')
                set_bg(al,'e8f1fc'); set_bg(av,'e8f1fc'); set_w(al,3.5); set_w(av,11.5)
                run(al.paragraphs[0],f'Phase {result.get("agent_phase",1)}',bold=True,size=9,color=ph_col)
                run(av.paragraphs[0],result['agent_label'],bold=True,size=10,color='1d1d1f')
                if result.get('agent_milestone'):
                    p_ms=av.add_paragraph(); run(p_ms,f'Milestone: {result["agent_milestone"]}',size=9,color='6e6e73')
                if result.get('orchestration'):
                    p_or=av.add_paragraph(); run(p_or,f'Pipeline: {result["orchestration"]}',size=9,color='1d1d1f',italic=True)

            doc.add_paragraph()


        # ── Roadmap ───────────────────────────────────────────────────────
        heading('Implementation Roadmap')
        for ph_title,ph_tasks2,col_hex2 in [
            ('Phase 1 — Quick Wins (0–3 months)',  [r for r in sorted_results if r['ai_readiness_score']>=70 and r.get('difficulty','').lower()=='easy'],   '34c759'),
            ('Phase 2 — Medium-Term (3–6 months)', [r for r in sorted_results if r['ai_readiness_score']>=50 and r.get('difficulty','').lower()=='medium'],  'ff9f0a'),
            ('Phase 3 — Advanced (6–12 months)',   [r for r in sorted_results if r['ai_readiness_score']>=40 and r.get('difficulty','').lower()=='hard'],    '0071e3'),
        ]:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10)
            run(p,ph_title,bold=True,size=12,color=col_hex2)
            if ph_tasks2:
                for r in ph_tasks2:
                    p=doc.add_paragraph(style='List Bullet')
                    run(p,r['task']['name'],bold=True,size=10,color='1d1d1f')
                    run(p,f"  /  {r['ai_readiness_score']:.0f}%  /  {r.get('estimated_hours_saved',0):.0f} hrs saved",size=9,color='6e6e73')
            else:
                p=doc.add_paragraph(); run(p,'No tasks in this phase.',size=9,color='6e6e73',italic=True)

        # ── Context sections ──────────────────────────────────────────────
        avg_human_edge = sum(r.get('human_edge_score') or 50 for r in results) / max(len(results),1)

        if context == 'individual':
            heading('Career Future Analysis', size=14, color='1d1d1f')

            # Countdown
            cw_tasks = [r for r in sorted_results if r.get('countdown_window')]
            if cw_tasks:
                p=doc.add_paragraph(); run(p,'Automation Countdown Clock',bold=True,size=12,color='1d1d1f')
                cw_map2={'now':'Automate NOW','12-24':'12–24 months','24-48':'24–48 months','48+':'48+ months'}
                cw_col2={'now':'ff3b30','12-24':'ff9f0a','24-48':'f5c518','48+':'34c759'}
                t=doc.add_table(rows=len(cw_tasks)+1,cols=3); t.style='Table Grid'
                for ci,hdr in enumerate(['Task','Risk Window','Human Edge']):
                    set_bg(t.rows[0].cells[ci],'f5f5f7')
                    run(t.rows[0].cells[ci].paragraphs[0],hdr,bold=True,size=9,color='6e6e73')
                for ri,r in enumerate(cw_tasks,1):
                    cw=r.get('countdown_window','24-48')
                    he=r.get('human_edge_score')
                    for ci,(val,col2) in enumerate([
                        (r['task']['name'],'1d1d1f'),
                        (cw_map2.get(cw,cw), cw_col2.get(cw,'6e6e73')),
                        (f'{he:.0f}/100' if he else '—','6e6e73'),
                    ]):
                        run(t.rows[ri].cells[ci].paragraphs[0],val,size=9,color=col2,bold=(ci==1))
                doc.add_paragraph()

            # Human Edge
            p=doc.add_paragraph(); run(p,'Human Edge Analysis',bold=True,size=12,color='1d1d1f')
            kv_table([
                ('AI Replacement Risk',    f'{score:.0f}%',            'ffe5e3'),
                ('Human Irreplaceability', f'{avg_human_edge:.0f}%',   'fff4e0'),
            ], col_widths=(5.0,10.0))
            insight2=('Strong human-essential components — amplify these while AI handles the rest.'
                      if avg_human_edge>=60 else 'Highly automatable role — pivot to higher human-edge functions now.')
            p=doc.add_paragraph(); run(p,insight2,size=9,color='6e6e73',italic=True)
            doc.add_paragraph()

            # Career Pivot
            all_skills2=[]; all_roles2=[]
            for r in sorted_results:
                try:
                    sk=_json.loads(r['pivot_skills']) if isinstance(r.get('pivot_skills'),str) else (r.get('pivot_skills') or [])
                    for s in (sk if isinstance(sk,list) else []):
                        if s not in all_skills2 and len(all_skills2)<6: all_skills2.append(s)
                except: pass
                try:
                    roles=_json.loads(r['pivot_roles']) if isinstance(r.get('pivot_roles'),str) else (r.get('pivot_roles') or [])
                    for role in (roles if isinstance(roles,list) else []):
                        if not any(x.get('role')==role.get('role') for x in all_roles2) and len(all_roles2)<5: all_roles2.append(role)
                except: pass

            if all_skills2 or all_roles2:
                p=doc.add_paragraph(); run(p,'Career Pivot Recommendations',bold=True,size=12,color='1d1d1f')
                if all_skills2:
                    p=doc.add_paragraph(); run(p,'Skills to Develop: ',bold=True,size=9,color='0071e3')
                    run(p,'  ·  '.join(all_skills2),size=9,color='1d1d1f')
                if all_roles2:
                    p=doc.add_paragraph(); run(p,'Adjacent Roles:',bold=True,size=9,color='0071e3')
                    t=doc.add_table(rows=len(all_roles2)+1,cols=3); t.style='Table Grid'
                    for ci,hdr in enumerate(['Role','Risk','Pivot Distance']):
                        set_bg(t.rows[0].cells[ci],'f5f5f7')
                        run(t.rows[0].cells[ci].paragraphs[0],hdr,bold=True,size=9,color='6e6e73')
                    for ri,role in enumerate(all_roles2,1):
                        rsk=role.get('risk','—')
                        rc_col='34c759' if rsk=='low' else ('ff9f0a' if rsk=='medium' else 'ff3b30')
                        run(t.rows[ri].cells[0].paragraphs[0],role.get('role','—'),size=9,color='1d1d1f')
                        run(t.rows[ri].cells[1].paragraphs[0],rsk,bold=True,size=9,color=rc_col)
                        run(t.rows[ri].cells[2].paragraphs[0],role.get('pivot_distance','—'),size=9,color='6e6e73')
                doc.add_paragraph()

        elif context == 'team':
            heading('Team Automation Strategy', size=14, color='1d1d1f')
            fte3=hours/1800
            p=doc.add_paragraph(); run(p,'Team Velocity Impact',bold=True,size=12,color='1d1d1f')
            kv_table([
                ('Hours freed / yr',  f'{hours:.0f}h',            'e8f1fc'),
                ('FTE equivalent',    f'{fte3:.1f} roles',         'e8f9ed'),
                ('Cost saved / yr',   f'\u20ac{savings:,.0f}',     'fff4e0'),
            ], col_widths=(5.0,10.0))
            doc.add_paragraph()

            p=doc.add_paragraph(); run(p,'Rollout Timeline',bold=True,size=12,color='1d1d1f')
            for ph_n,ph_t,ph_c in [('Phase 1 — Quick Wins (0–3mo)',[r for r in sorted_results if r.get('difficulty','').lower()=='easy'],'34c759'),
                                    ('Phase 2 — Medium (3–12mo)',[r for r in sorted_results if r.get('difficulty','').lower()=='medium'],'ff9f0a'),
                                    ('Phase 3 — Strategic (12–36mo)',[r for r in sorted_results if r.get('difficulty','').lower()=='hard'],'0071e3')]:
                ph_h=sum(r.get('estimated_hours_saved',0) for r in ph_t)
                p=doc.add_paragraph(); run(p,f'{ph_n}  ',bold=True,size=10,color=ph_c)
                run(p,f'{ph_h:.0f}h/yr · {len(ph_t)} tasks',size=9,color='6e6e73')
            doc.add_paragraph()

            sprint=[r for r in sorted_results if r.get('difficulty','').lower()=='easy'][:5]
            p=doc.add_paragraph(); run(p,'90-Day Sprint Plan',bold=True,size=12,color='1d1d1f')
            if sprint:
                t=doc.add_table(rows=len(sprint)+1,cols=4); t.style='Table Grid'
                for ci,hdr in enumerate(['#','Task','Score','Hours/yr']):
                    set_bg(t.rows[0].cells[ci],'f5f5f7')
                    run(t.rows[0].cells[ci].paragraphs[0],hdr,bold=True,size=9,color='6e6e73')
                for ri,r in enumerate(sprint,1):
                    run(t.rows[ri].cells[0].paragraphs[0],str(ri),bold=True,size=9,color='34c759')
                    run(t.rows[ri].cells[1].paragraphs[0],r['task']['name'],size=9,color='1d1d1f')
                    run(t.rows[ri].cells[2].paragraphs[0],f'{r["ai_readiness_score"]:.0f}%',bold=True,size=9,color='0071e3')
                    run(t.rows[ri].cells[3].paragraphs[0],f'{r.get("estimated_hours_saved",0):.0f}h',size=9,color='1d1d1f')
            else:
                p=doc.add_paragraph(); run(p,'No easy-difficulty tasks — focus on Phase 2.',size=9,color='6e6e73',italic=True)
            doc.add_paragraph()

        elif context == 'company':
            heading('Strategic Business Analysis', size=14, color='1d1d1f')

            p=doc.add_paragraph(); run(p,'AI-First Competitor Gap',bold=True,size=12,color='1d1d1f')
            kv_table([
                ('If you automate now',      f'\u20ac{savings:,.0f}/yr',         'e8f9ed'),
                ('If you wait 12 months',    f'\u20ac{savings*0.35:,.0f}/yr',    'fff4e0'),
                ('AI-first competitor edge', f'\u20ac{savings*1.4:,.0f}/yr',     'ffe5e3'),
            ], col_widths=(5.5,9.5))
            doc.add_paragraph()

            fte4=hours/1800
            p=doc.add_paragraph(); run(p,'Headcount Signal',bold=True,size=12,color='1d1d1f')
            kv_table([
                ('Hours freed / yr',  f'{hours:.0f}h',                          'e8f1fc'),
                ('FTE equivalent',    f'{fte4:.1f} roles',                       'f5f5f7'),
                ('Saved per FTE',     f'\u20ac{savings/max(fte4,0.1):,.0f}',    'f5f5f7'),
            ], col_widths=(5.0,10.0))
            doc.add_paragraph()

            gap2=max(0,81-round(score))
            p=doc.add_paragraph(); run(p,'Industry Benchmark',bold=True,size=12,color='1d1d1f')
            kv_table([
                ('Your score',     f'{score:.0f}%',  'e8f1fc'),
                ('Sector average', '58%',             'f5f5f7'),
                ('Top 10% orgs',   '81%',             'e8f9ed'),
                ('Gap to close',   f'{gap2}%',        'fff4e0'),
            ], col_widths=(5.0,10.0))
            doc.add_paragraph()

            # Board summary
            quick2=len([r for r in results if r.get('difficulty','').lower()=='easy'])
            flags2=len([r for r in results if r.get('risk_level') != 'safe'])
            p=doc.add_paragraph(); run(p,'Board-Ready Executive Summary',bold=True,size=12,color='1d1d1f')
            bs=doc.add_table(rows=9,cols=2); bs.style='Table Grid'
            board=[('Workflow',workflow.get('name','')),('Industry',workflow.get('industry','General')),
                   ('Automation potential',f'{score:.0f}% of tasks'),('Annual savings',f'\u20ac{savings:,.0f}'),
                   ('Hours reclaimed',f'{hours:.0f}h/yr'),('FTE equivalent',f'{fte4:.1f} roles'),
                   ('Quick wins (90 days)',f'{quick2} tasks'),('Risk flags',f'{flags2} require review'),
                   ('Recommendation','Begin 90-day sprint. Prioritise quick wins. Redeploy freed capacity.')]
            for ri,(k,v) in enumerate(board):
                set_bg(bs.rows[ri].cells[0],'1d1d1f'); set_bg(bs.rows[ri].cells[1],'2c2c2e')
                set_w(bs.rows[ri].cells[0],5.5); set_w(bs.rows[ri].cells[1],9.5)
                run(bs.rows[ri].cells[0].paragraphs[0],k,bold=True,size=9,color='86868b')
                run(bs.rows[ri].cells[1].paragraphs[0],v,size=10,color='ffffff')
            doc.add_paragraph()

        # ── Conclusion ────────────────────────────────────────────────────
        heading('Conclusion')
        p=doc.add_paragraph()
        run(p,f'This analysis identified automation opportunities across {len(results)} tasks in {workflow["name"]}. ',size=10,color='1d1d1f')
        run(p,f'Implementing recommendations could save {hours:.0f} hours annually',bold=True,size=10,color='1d1d1f')
        run(p,f', worth approximately \u20ac{savings:,.0f}.',size=10,color='1d1d1f')
        for step in ['Start with Phase 1 quick wins for immediate ROI within weeks.',
                     'Apply human-in-the-loop for tasks scored 40–70 to manage risk.',
                     'Track time saved monthly to measure and communicate impact.',
                     'Revisit quarterly — new AI tools appear constantly.',
                     'Train team members on newly automated workflows.']:
            p=doc.add_paragraph(style='List Bullet'); run(p,step,size=10,color='1d1d1f')
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p,'Generated by WorkScanAI — AI-Powered Workflow Analysis',size=8,color='6e6e73')
        doc.save(output_path)
        return output_path


    # ─────────────────────────────────────────────────────────────────────────
    # COMBINED REPORTS
    # ─────────────────────────────────────────────────────────────────────────

    @staticmethod
    def generate_combined_pdf_report(analyses_list: List[Dict], output_path: str):
        """One PDF — master cover + each workflow as a full section."""
        import tempfile, os
        from reportlab.platypus import SimpleDocTemplate

        doc = SimpleDocTemplate(output_path, pagesize=A4,
            leftMargin=18*mm, rightMargin=18*mm, topMargin=20*mm, bottomMargin=24*mm)
        W = A4[0] - 36*mm
        story = []
        s = getSampleStyleSheet()

        def style(name, parent='Normal', **kw):
            return ParagraphStyle(name, parent=s[parent], **kw)

        ST = {
            'section_title': style('st2',fontSize=18,leading=22,textColor=GRAY_900,
                                    spaceBefore=18,spaceAfter=8,fontName='Helvetica-Bold'),
            'label': style('lb2',fontSize=9,leading=12,textColor=GRAY_600,fontName='Helvetica-Bold',spaceAfter=2),
            'body':  style('bo2',fontSize=10,leading=15,textColor=GRAY_900,fontName='Helvetica',spaceAfter=6),
            'rec':   style('rc2',fontSize=10,leading=15,textColor=GRAY_900,fontName='Helvetica',leftIndent=10,spaceAfter=8),
        }

        total_savings = sum(a['annual_savings'] for a in analyses_list)
        total_hours   = sum(a['hours_saved']    for a in analyses_list)
        total_tasks   = sum(len(a['results'])   for a in analyses_list)

        # Master cover
        story.append(Table([['']],colWidths=[W],rowHeights=[4],
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE)])))
        story.append(Spacer(1,16*mm))
        story.append(Paragraph('WorkScanAI',style('brand2',fontSize=11,textColor=BLUE,fontName='Helvetica-Bold',spaceAfter=10)))
        story.append(Paragraph('Combined Workflow Automation\nAnalysis Report',
            style('ct2',fontSize=28,leading=34,textColor=GRAY_900,spaceAfter=6,fontName='Helvetica-Bold')))
        story.append(Paragraph(f'{len(analyses_list)} Workflows',
            style('wn2',fontSize=20,leading=24,textColor=BLUE,fontName='Helvetica-Bold',spaceAfter=10)))
        story.append(Paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')}",
            style('cm2',fontSize=10,leading=14,textColor=GRAY_600,fontName='Helvetica')))
        story.append(Spacer(1,14*mm))

        story.append(Table([[
            Paragraph(f'<font size="36"><b>{len(analyses_list)}</b></font><br/><font size="12" color="#6e6e73">workflows</font>',
                style('hs2',fontName='Helvetica-Bold',alignment=TA_CENTER,leading=42)),
            Table([[Paragraph('<b>Total Savings</b>',ST['label']),''],
                   [Paragraph(f'\u20ac{total_savings:,.0f}',style('sv2',fontSize=20,fontName='Helvetica-Bold',textColor=GRAY_900)),''],
                   [Paragraph(f'{total_hours:.0f} hours reclaimed',style('hr2',fontSize=10,textColor=GRAY_600,fontName='Helvetica')),'']],
                  colWidths=[W*0.38,W*0.05]),
            Table([[Paragraph('<b>Total Tasks</b>',ST['label']),''],
                   [Paragraph(str(total_tasks),style('tc2',fontSize=20,fontName='Helvetica-Bold',textColor=GRAY_900)),''],
                   [Paragraph('analyzed',style('tl2',fontSize=10,textColor=GRAY_600,fontName='Helvetica')),'']],
                  colWidths=[W*0.25,W*0.05]),
        ]],colWidths=[W*0.27,W*0.44,W*0.29],
        style=TableStyle([('BACKGROUND',(0,0),(-1,-1),BLUE_LIGHT),
            ('TOPPADDING',(0,0),(-1,-1),16),('BOTTOMPADDING',(0,0),(-1,-1),16),
            ('LEFTPADDING',(0,0),(0,-1),20),('LEFTPADDING',(1,0),(1,-1),10),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEAFTER',(0,0),(0,-1),1,BLUE)])))
        story.append(Spacer(1,10*mm))

        # Index
        idx_rows=[[Paragraph('<b>#</b>',ST['label']),Paragraph('<b>Workflow</b>',ST['label']),
                   Paragraph('<b>Score</b>',ST['label']),Paragraph('<b>Savings</b>',ST['label'])]]
        for i,a in enumerate(analyses_list,1):
            sc,_=score_color(a['automation_score'])
            idx_rows.append([
                Paragraph(str(i),style(f'ix{i}',fontSize=10,fontName='Helvetica',textColor=GRAY_600)),
                Paragraph(a['workflow']['name'],style(f'iw{i}',fontSize=10,fontName='Helvetica-Bold',textColor=GRAY_900)),
                Paragraph(f"{a['automation_score']:.0f}%",style(f'is{i}',fontSize=10,fontName='Helvetica-Bold',textColor=sc)),
                Paragraph(f"\u20ac{a['annual_savings']:,.0f}",style(f'iv{i}',fontSize=10,fontName='Helvetica',textColor=GRAY_900)),
            ])
        story.append(Table(idx_rows,colWidths=[10*mm,W-10*mm-30*mm-40*mm,30*mm,40*mm],
            style=TableStyle([('BACKGROUND',(0,0),(-1,0),GRAY_100),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,GRAY_100]),
                ('TOPPADDING',(0,0),(-1,-1),7),('BOTTOMPADDING',(0,0),(-1,-1),7),
                ('LEFTPADDING',(0,0),(-1,-1),10),('LINEBELOW',(0,0),(-1,-1),0.3,GRAY_200)])))
        story.append(PageBreak())

        # Per-workflow: full sections
        for w_idx, analysis_data in enumerate(analyses_list):
            wf = analysis_data['workflow']
            sc, sc_light = score_color(analysis_data['automation_score'])
            story.append(Table([['']],colWidths=[W],rowHeights=[3],
                style=TableStyle([('BACKGROUND',(0,0),(-1,-1),sc)])))
            story.append(Spacer(1,10*mm))
            story.append(Paragraph(f'Workflow {w_idx+1} of {len(analyses_list)}',
                style(f'wl{w_idx}',fontSize=10,textColor=GRAY_600,fontName='Helvetica',spaceAfter=6)))
            story.append(Paragraph(wf['name'],
                style(f'wt{w_idx}',fontSize=24,leading=28,textColor=GRAY_900,fontName='Helvetica-Bold',spaceAfter=8)))

            # Mini KPIs
            sc2,sc2_light=score_color(analysis_data['automation_score'])
            story.append(Table([[
                Paragraph(f'<b>{analysis_data["automation_score"]:.0f}%</b><br/><font size="9" color="#6e6e73">Score</font>',
                    style(f'ws{w_idx}',fontSize=20,fontName='Helvetica-Bold',textColor=sc2,alignment=TA_CENTER,leading=26)),
                Paragraph(f'<b>\u20ac{analysis_data["annual_savings"]:,.0f}</b><br/><font size="9" color="#6e6e73">Savings</font>',
                    style(f'wv{w_idx}',fontSize=16,fontName='Helvetica-Bold',textColor=GRAY_900,alignment=TA_CENTER,leading=22)),
                Paragraph(f'<b>{analysis_data["hours_saved"]:.0f}h</b><br/><font size="9" color="#6e6e73">Hrs/yr</font>',
                    style(f'wh{w_idx}',fontSize=16,fontName='Helvetica-Bold',textColor=GRAY_900,alignment=TA_CENTER,leading=22)),
                Paragraph(f'<b>{len(analysis_data["results"])}</b><br/><font size="9" color="#6e6e73">Tasks</font>',
                    style(f'wc{w_idx}',fontSize=16,fontName='Helvetica-Bold',textColor=GRAY_900,alignment=TA_CENTER,leading=22)),
            ]],colWidths=[W/4]*4,
            style=TableStyle([('BACKGROUND',(0,0),(-1,-1),sc2_light),
                ('TOPPADDING',(0,0),(-1,-1),12),('BOTTOMPADDING',(0,0),(-1,-1),12),
                ('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LINEAFTER',(0,0),(2,-1),0.5,sc2)])))
            story.append(Spacer(1,8*mm))

            # Full task blocks with all features
            story.append(Paragraph('Task Analysis',ST['section_title']))
            story.append(HRFlowable(width=W,thickness=0.5,color=GRAY_200,spaceAfter=8))
            sorted_wf = sorted(analysis_data['results'], key=lambda x: x['ai_readiness_score'], reverse=True)
            for blk in ReportGenerator._pdf_task_blocks(sorted_wf, analysis_data, W, s, ST, style):
                story.append(blk)

            # Context sections
            ReportGenerator._pdf_context_sections(story, analysis_data, W, style, ST)

            if w_idx < len(analyses_list) - 1:
                story.append(PageBreak())

        # Combined conclusion
        story.append(PageBreak())
        story.append(Paragraph('Combined Summary',ST['section_title']))
        story.append(HRFlowable(width=W,thickness=0.5,color=GRAY_200,spaceAfter=10))
        story.append(Paragraph(
            f'This combined report covers <b>{len(analyses_list)} workflows</b> with '
            f'<b>{total_tasks} total tasks</b>. Implementing all recommendations could save '
            f'<b>{total_hours:.0f} hours annually</b>, worth approximately <b>\u20ac{total_savings:,.0f}</b>.',
            style('conc',fontSize=10,leading=15,textColor=GRAY_900,fontName='Helvetica',spaceAfter=6)))

        doc.build(story, canvasmaker=NumberedCanvas)
        return output_path

    @staticmethod
    def generate_combined_docx_report(analyses_list: List[Dict], output_path: str):
        """One DOCX — for each workflow, generate a full DOCX and merge paragraphs."""
        if Document is None:
            raise ImportError("python-docx not installed")
        import tempfile, os

        combined = Document()
        for sec in combined.sections:
            sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
            sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

        def add_run(para, text, bold=False, size=None, color=None, italic=False):
            r=para.add_run(text); r.bold=bold; r.italic=italic
            if size: r.font.size=Pt(size)
            if color: r.font.color.rgb=RGBColor(*tuple(int(color[i:i+2],16) for i in (0,2,4)))
            return r

        total_savings=sum(a['annual_savings'] for a in analyses_list)
        total_tasks=sum(len(a['results']) for a in analyses_list)

        # Master cover
        p=combined.add_paragraph(); add_run(p,'WorkScanAI',bold=True,size=11,color='0071e3')
        p=combined.add_paragraph(); add_run(p,'Combined Workflow Automation Analysis Report',bold=True,size=26,color='1d1d1f')
        p=combined.add_paragraph(); add_run(p,f'{len(analyses_list)} Workflows  ·  {total_tasks} Tasks  ·  \u20ac{total_savings:,.0f} Potential Savings',size=13,color='6e6e73')
        p=combined.add_paragraph(); add_run(p,f"Generated {datetime.now().strftime('%B %d, %Y')}",size=9,color='6e6e73')
        p.paragraph_format.space_after=Pt(14)

        # Write each workflow as its own DOCX, then copy elements
        for w_idx, analysis_data in enumerate(analyses_list):
            tmp_path = os.path.join(tempfile.gettempdir(), f'_wf_tmp_{w_idx}.docx')
            ReportGenerator.generate_docx_report(analysis_data, tmp_path)
            src = Document(tmp_path)

            # Divider
            p=combined.add_paragraph(); p.paragraph_format.page_break_before=True
            add_run(p,f'WORKFLOW {w_idx+1} OF {len(analyses_list)}',bold=True,size=9,color='6e6e73')

            # Copy all body elements
            from docx.oxml import OxmlElement as OE
            from copy import deepcopy
            for element in src.element.body:
                combined.element.body.append(deepcopy(element))

            os.remove(tmp_path)

        combined.save(output_path)
        return output_path
